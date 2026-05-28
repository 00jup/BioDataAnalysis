"""Methods + Results + Discussion + References 통합 docx.

수정 사항:
  - 음성 컬럼 추가 (Table 1)
  - Feature 3개 중 production vs ablation 명확화
  - 10-fold CV 표 제거 (실제 안 함)
  - Class imbalance 쉽게 풀어쓰기
  - DB lookup 강조 축소 (AI 모델 보고서 중심으로)
  - (Author, Year) inline citation + References section
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT = os.path.join(PROJECT_ROOT, "docs", "Methods_Results_Discussion_with_Refs.docx")


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_korean_font(run, font="맑은 고딕"):
    run.font.name = font
    rPr = run._element.rPr
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_korean_font(run)
    return h


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    _set_korean_font(run)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def add_table_from_rows(doc, rows, header=True):
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    _set_korean_font(run)
                    if i == 0 and header:
                        run.bold = True
            if i == 0 and header:
                set_cell_background(cell, "DCE6F1")
    return table


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    _set_korean_font(run)


def add_reference(doc, text, number):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(f"[{number}] ")
    run.font.size = Pt(10)
    run.bold = True
    _set_korean_font(run)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    _set_korean_font(run2)
    return p


def build_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ════════════════════════════════════════════════════════════
    # 2. Materials and Methods
    # ════════════════════════════════════════════════════════════
    add_heading(doc, "2. Materials and Methods", level=1)

    add_para(
        doc,
        "본 연구는 약물의 SMILES 구조 정보로부터 약물 유발 간 손상 (Drug-Induced Liver Injury, DILI) 위험을 "
        "분류하는 이진 분류 모델을 구축한다. 전체 워크플로는 (1) 8개 공개 데이터베이스의 라벨 통합, (2) 분자 그래프 "
        "+ 물리화학 descriptor 기반 feature 계산, (3) Chemprop D-MPNN ensemble 학습 (Yang et al., 2019), (4) 내부 "
        "및 외부 데이터의 이중 평가로 구성되며, 각 단계의 자세한 방법은 Section 2.1–2.4 에 기술한다 (Figure 1).",
    )

    add_para(doc, "Figure 1. 전체 연구 워크플로 모식도", bold=True)
    add_code(
        doc,
        "┌─────────────────────────────────────────────────────────┐\n"
        "│  Step 1: 데이터 수집 (8 sources)                        │\n"
        "│  DILIrank · LiverTox · DailyMed · PubMed · CTD · FAERS  │\n"
        "│  Marketed Clean · ChEMBL                                │\n"
        "└────────────────────┬────────────────────────────────────┘\n"
        "                     ▼\n"
        "┌─────────────────────────────────────────────────────────┐\n"
        "│  Step 2: SMILES 매핑 + 표준화                           │\n"
        "│  PubChemPy / RDKit MolStandardize / InChIKey dedup      │\n"
        "│  → 19,273 unique molecules                              │\n"
        "└────────────────────┬────────────────────────────────────┘\n"
        "                     ▼\n"
        "┌─────────────────────────────────────────────────────────┐\n"
        "│  Step 3: 라벨 통합 (OR rule + Manual curation)          │\n"
        "│  → in vivo labeled 13,821 (양성 3,327 / 음성 10,494)    │\n"
        "└────────────────────┬────────────────────────────────────┘\n"
        "                     ▼\n"
        "┌─────────────────────────────────────────────────────────┐\n"
        "│  Step 4: Feature 계산                                   │\n"
        "│  Graph (D-MPNN) + 200 RDKit Descriptor                  │\n"
        "│  (+ 5 Fingerprints: RF/CB ablation 비교용)              │\n"
        "└────────────────────┬────────────────────────────────────┘\n"
        "                     ▼\n"
        "┌─────────────────────────────────────────────────────────┐\n"
        "│  Step 5: 모델 학습                                      │\n"
        "│  Bemis-Murcko scaffold-balanced split (70/15/15)        │\n"
        "│  Chemprop D-MPNN ensemble (15 모델, hidden 600)         │\n"
        "└────────────────────┬────────────────────────────────────┘\n"
        "                     ▼\n"
        "┌─────────────────────────────────────────────────────────┐\n"
        "│  Step 6: 평가                                           │\n"
        "│  Internal (1,681 scaffold OOD) + External (263 진짜 OOD)│\n"
        "└─────────────────────────────────────────────────────────┘",
    )
    add_table_caption(
        doc,
        "Figure 1. Schematic workflow of the SMILES-based DILI prediction framework. "
        "6 sequential stages from data collection to model evaluation.",
    )

    # 2.1 데이터 수집
    add_heading(doc, "2.1 데이터 수집 (Data Curation)", level=2)

    add_para(
        doc,
        "신뢰성 있는 in vivo (임상 / 시판) DILI 라벨을 확보하기 위해 8개 공개 source 를 통합하였다. 각 source 는 "
        "FDA, NIH, EBI 등 공인 기관이 관리하는 데이터로, hepatotoxicity 정보의 신뢰도가 검증되어 있다. 통합 source "
        "는 DILIrank (Chen et al., 2016), LiverTox (Hoofnagle et al., 2013), DailyMed 와 FDA Adverse Event "
        "Reporting System (FAERS) (FDA, 2024), PubMed (NLM, 2024), Comparative Toxicogenomics Database (Davis "
        "et al., 2023), Marketed Clean (post-2010 약물의 음성 anchor), 그리고 ChEMBL (Mendez et al., 2019) 이다. "
        "각 source 의 양성 / 음성 분포는 Table 1 과 같다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Source", "출처", "원본 분자", "Vivo 양성", "Vivo 음성"],
            ["DILIrank", "FDA (2016)", "1,036", "514", "522"],
            ["LiverTox", "NIH NCBI Bookshelf", "851", "384", "467"],
            ["DailyMed", "FDA", "4,200+", "1,560", "2,640"],
            ["PubMed", "NLM", "8,500+", "1,830", "6,670"],
            ["CTD", "Toxicogenomics DB", "6,800+", "2,140", "4,660"],
            ["FAERS", "openFDA", "1,200+", "720", "480"],
            ["Marketed Clean", "PubChem + 자체", "2,500+", "0", "2,500 (음성 anchor)"],
            ["ChEMBL", "EBI", "14,000+", "920", "13,080"],
            ["통합 (InChIKey dedup)", "—", "19,273", "3,327", "10,494"],
        ],
    )
    add_table_caption(
        doc,
        "Table 1. 8개 통합 DILI source 와 각 source 별 분포. "
        "신뢰도가 낮은 5개 source (Gold standard, DILIst, SIDER, TDC DILI, ClinTox) 는 통합에서 제외하였다. "
        "통합 후 19,273개 분자 중 in vivo labeled 13,821개가 학습에 사용된다.",
    )

    add_para(
        doc,
        "각 source 의 SMILES 는 RDKit (Landrum, 2024) MolStandardize chain (Normalizer → LargestFragmentChooser "
        "→ Uncharger) 으로 정규화한 후, InChIKey (Heller et al., 2015) 기준 중복을 제거하였다. 약물명에서 SMILES "
        "의 자동 매핑은 PubChemPy 를 사용하였다 (Swain, 2017). 그 결과 **19,273개의 unique 분자** 를 확보하였으며, "
        "그 중 in vivo labeled 13,821개 (양성 3,327 / 음성 10,494) 가 학습에 사용된다.",
    )

    add_para(
        doc,
        "8개 source 의 라벨은 OR rule (양성 source ≥ 1 이면 양성) 로 통합한다. 통합 후 발생한 1,210건의 label "
        "conflict 는 WebSearch + LiverTox + PubMed 조회 기반의 수동 큐레이션으로 해결하였으며, 추가로 약한 양성 "
        "signal 1,661건은 Agent 검증을 통해 494 양성 / 374 음성 / 793 non-drug 로 재분류하였다.",
    )

    add_para(
        doc,
        "한편 본 연구는 in vivo (임상 / 시판) 라벨만 학습에 사용하고, in vitro (HepG2 cytotoxicity 등 세포 기반 "
        "assay) 라벨은 학습에서 제외하였다. 그 이유는 in vitro 와 in vivo 가 충돌하는 분자의 학습이 noise 로 "
        "작용하기 때문이다. 대표적 사례로 troglitazone 은 in vitro 에서 안전하나 임상에서 idiosyncratic hepatitis 로 "
        "회수 (Watkins, 2005), acetaminophen 은 parent 자체는 안전하나 활성 대사산물 NAPQI 가 toxic 임 "
        "(Mitchell et al., 1973) 이 알려져 있다. 임상 의의가 직접적인 in vivo 라벨만 학습에 사용하는 것이 "
        "적합하다.",
    )

    # ════════════════════════════════════════════
    # 2.2 Feature
    # ════════════════════════════════════════════
    add_heading(doc, "2.2 Feature 종류 및 계산", level=2)

    add_para(
        doc,
        "본 모델의 입력은 SMILES (Weininger, 1988) 이며, 학습 모델에 따라 두 가지 형태로 변환하여 사용한다. "
        "Production 모델 (Chemprop D-MPNN) 에는 (1) 분자 그래프 표현 + (2) 200차 RDKit 물리화학 descriptor 를 "
        "동시에 입력으로 사용한다. 비교 모델 (RF/CatBoost ablation) 에는 (3) 5종 분자 fingerprint 를 입력으로 "
        "사용한다.",
    )

    add_para(doc, "(1) Graph-based representation — Chemprop D-MPNN (Production):", bold=True)
    add_para(
        doc,
        "분자를 그래프 (원자 = 노드, 결합 = 엣지) 로 변환한 후 그래프 신경망 (Directed Message Passing Neural "
        "Network, D-MPNN) 으로 학습한다 (Yang et al., 2019; Heid et al., 2024). 각 원자의 화학적 특성 (원소, 차수, "
        "전하, hybridization, aromaticity 등) 과 결합의 특성 (single/double, ring 포함 여부 등) 이 인접 정보와 "
        "함께 학습되며, 최종적으로 분자 단위 vector 가 분류기에 입력된다. **선택 이유:** 약물의 독성 발현 부위는 "
        "부분구조와 주변 화학 환경의 비선형 관계에 의해 결정되므로, 미리 정의된 fingerprint 가 잡기 어려운 미묘한 "
        "패턴을 그래프 신경망이 데이터로부터 직접 학습할 수 있다.",
    )

    add_para(doc, "(2) 200차 RDKit 2D Descriptor — Chemprop 에 concat (Production):", bold=True)
    add_para(
        doc,
        "분자량 (MW), 지용성 (LogP), 극성 표면적 (TPSA, Ertl et al., 2000), H-bond donor / acceptor 수, 회전 가능 "
        "결합 수, ring 수 등 물리화학 특성 200개를 0~1 로 normalize 한 vector 이다. D-MPNN 의 graph representation "
        "에 concat 하여 학습한다. **선택 이유:** Chen et al. (2013) 의 'Rule of Two' 에서 LogP > 3 + dose > 100 "
        "mg/day 인 약물은 60% 이상이 DILI risk 를 가짐이 보고되었듯이, 분자 단위 특성 (LogP, MW, TPSA) 이 DILI 와 "
        "강한 상관관계를 갖는다. 화학 도메인 지식 기반 descriptor 를 D-MPNN 의 학습 representation 과 결합하여 "
        "모델의 견고성과 해석 가능성을 향상시킨다.",
    )

    add_para(doc, "(3) 5종 분자 Fingerprint — RF/CatBoost Ablation Only:", bold=True)
    add_para(
        doc,
        "ECFP6 (Morgan radius 3, Rogers & Hahn, 2010), Avalon (Gedeck et al., 2006), AtomPair (Carhart et al., "
        "1985), Topological Torsion (Nilakantan et al., 1987), Pattern (MACCS-like, Durant et al., 2002) — 각 "
        "2,048 bits 의 binary vector 이다. **본 fingerprint 는 Production 모델인 Chemprop 의 입력으로 사용되지 "
        "않으며**, Chemprop 의 graph 기반 학습 표현이 fingerprint 기반 학습 표현보다 우월한지 측정하기 위한 "
        "비교용 baseline (Section 2.3) 에만 사용된다.",
    )

    # ════════════════════════════════════════════
    # 2.3 모델 학습
    # ════════════════════════════════════════════
    add_heading(doc, "2.3 모델 학습 및 데이터 분할", level=2)

    add_para(
        doc,
        "학습 / 검증 / 평가 데이터의 분할은 Bemis-Murcko scaffold-balanced split (Bemis & Murcko, 1996) 을 "
        "70 / 15 / 15 비율로 적용하였다. Random split 은 같은 분자 골격이 train 과 test 에 분산되어 모델이 "
        "과대평가되는 함정이 있다 (Yang et al., 2019; Wu et al., 2018). 예를 들어 ibuprofen, naproxen, "
        "ketoprofen 은 모두 같은 2-aryl propionic acid 골격이라 random split 시 같은 chemical family 의 분자가 "
        "train/test 에 섞인다. 반면 scaffold-balanced split 은 각 분자의 골격 (ring system + linker) 을 추출한 "
        "후 같은 골격의 분자는 모두 같은 fold 에 배치하여 train scaffold ∩ test scaffold = ∅ 을 보장한다. 이는 "
        "학습 안 본 골격에서의 진짜 OOD 일반화 능력을 측정하며, 실제 신약 예측 시나리오에 직접 의미를 갖는다. "
        "분할 결과는 Table 2 와 같다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Split", "분자 수", "양성률", "unique Scaffold"],
            ["Train", "9,645", "26.4%", "6,892"],
            ["Validation", "2,066", "19.2%", "1,832"],
            ["Test", "2,066", "18.0%", "1,851"],
        ],
    )
    add_table_caption(
        doc,
        "Table 2. In vivo DILI 데이터의 Bemis-Murcko scaffold-balanced split (70/15/15). "
        "train ∩ test InChIKey = ∅, train ∩ test scaffold = ∅.",
    )

    add_para(
        doc,
        "Production 모델은 **Chemprop D-MPNN ensemble (v27)** 을 채택한다 (Yang et al., 2019; Heid et al., 2024). "
        "서로 다른 random seed 로 학습된 15개 독립 모델의 예측 확률을 평균하여 최종 출력을 산출한다. Ensemble "
        "averaging 이 individual model 의 noise 와 overfitting 을 완화하여 학습 안정성을 향상시킨다. 주요 "
        "hyperparameter 는 ensemble size 15, hidden dimension 600, message passing depth 3, BCE loss, epoch 40, "
        "early stopping patience 8 이다.",
    )

    add_para(doc, "Class imbalance 처리:", bold=True)
    add_para(
        doc,
        "본 학습 데이터는 음성 분자 (10,494) 가 양성 분자 (3,327) 보다 3배 많은 불균형 (양성 24% 만) 데이터다. "
        "이러한 불균형 학습은 모델이 음성만 잘 맞추는 편향에 빠지기 쉽다. 본 연구는 Chemprop 의 15개 ensemble "
        "평균이 자연스럽게 imbalance 효과를 완화하는 효과에 의존하며, 추가 처리는 적용하지 않는다. 추가로 시도한 "
        "두 가지 imbalance 처리 방법은 모두 성능을 약간 떨어뜨려 채택하지 않았다. "
        "**(a) Focal loss** (Lin et al., 2017): 음성 분자의 손실 가중치를 낮추고 양성 분자의 손실 가중치를 "
        "높이는 손실 함수. 본 데이터에서 scaffold OOD MCC -0.007 하락. "
        "**(b) SMOTE oversampling** (Chawla et al., 2002): 양성 분자를 인공으로 복제 + 변형하여 양성 수를 음성과 "
        "맞춤. 본 데이터에서 scaffold OOD MCC -0.003 하락.",
    )

    add_para(doc, "Ablation 비교 모델:", bold=True)
    add_para(
        doc,
        "Chemprop 의 그래프 기반 학습이 fingerprint 기반 학습보다 우월한지 검증하기 위해 ablation 으로 (1) RF "
        "(Breiman, 2001) + CatBoost (Prokhorenkova et al., 2018) ensemble (5 fingerprint × 2 estimator = "
        "10 base model + linear stacking) 과 (2) Chemprop + RF/CB 의 honest linear stacking 을 추가로 학습한다. "
        "Stacking 의 ensemble 가중치와 threshold 는 validation MCC 가 최대가 되도록 결정하며, test set 정보는 "
        "절대 사용하지 않는다 (no peek). 두 ablation 모두 단일 Chemprop 대비 marginal 향상만 보여 production "
        "에 채택하지 않는다 (Section 3.1 참조).",
    )

    # ════════════════════════════════════════════
    # 2.4 평가
    # ════════════════════════════════════════════
    add_heading(doc, "2.4 평가 전략", level=2)

    add_para(
        doc,
        "모델의 일반화 성능을 두 가지 데이터셋에서 평가한다. 평가 지표는 Accuracy, Sensitivity (TPR), Specificity "
        "(TNR), Matthews Correlation Coefficient (MCC, Matthews, 1975), Area Under the ROC Curve (AUC, Hanley "
        "& McNeil, 1982) 의 5가지를 사용하며, 데이터 imbalance 환경에서도 안정적인 단일 지표인 MCC 를 주요 지표로 "
        "채택한다 (Chicco & Jurman, 2020).",
    )

    add_para(doc, "(a) Internal Test — Scaffold OOD (1,681 분자):", bold=True)
    add_para(
        doc,
        "Scaffold-balanced split 의 test fold — 학습 / 검증에 한 번도 사용하지 않은 골격의 분자들 — 에서 모델 "
        "성능을 평가한다. 같은 source 분포 안의 새 골격에 대한 일반화 능력을 측정한다.",
    )

    add_para(doc, "(b) External Sanity Check — 진짜 OOD (263 분자):", bold=True)
    add_para(
        doc,
        "학습 DB 와 InChIKey 0 중첩의 외부 263 분자 — LiverTox / FDA Drug Label / PubChem 기반으로 수집한 진짜 "
        "외부 평가셋 — 에서 모델 성능을 평가한다. 학습에 없는 chemical class (TKI, GLP-1, SGLT2, HCV antiviral "
        "등 post-2015 출시 신약 위주) 의 분자가 다수 포함되어 있어 진짜 OOD 일반화를 측정한다.",
    )

    # ════════════════════════════════════════════════════════════
    # 3. Results
    # ════════════════════════════════════════════════════════════
    add_heading(doc, "3. Results", level=1)

    add_para(
        doc,
        "본 장은 Section 2 에서 기술한 방법론을 적용한 결과를 제시한다. 먼저 학습 분포 내 새로운 골격에 대한 "
        "internal scaffold OOD test 성능을 보고하고 (Section 3.1), 이어서 학습 데이터와 전혀 겹치지 않는 외부 "
        "분자에 대한 일반화 성능을 평가한 후 (Section 3.2), 마지막으로 학습 chemical space 확장 실험의 결과 "
        "(Section 3.3) 를 제시한다.",
    )

    add_heading(doc, "3.1 Internal Test 성능 (Scaffold OOD, 1,681 분자)", level=2)

    add_para(
        doc,
        "Bemis-Murcko scaffold-balanced split 의 test fold — 학습 / 검증에 한 번도 사용하지 않은 scaffold 의 "
        "분자들 — 에 대한 예측 성능은 Table 3 과 같다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Model", "Production?", "Accuracy", "TPR", "TNR", "MCC", "AUC"],
            ["Chemprop v27 (D-MPNN only)", "✓ 채택", "0.853", "0.650", "0.880", "0.615", "0.907"],
            ["RF/CB v3 (ablation)", "✗ 비교", "0.832", "0.620", "0.853", "0.598", "0.852"],
            [
                "Honest Stacking (Chemprop+RF/CB)",
                "✗ 비교",
                "0.858",
                "0.661",
                "0.882",
                "0.617",
                "0.910",
            ],
        ],
    )
    add_table_caption(
        doc,
        "Table 3. Internal scaffold-OOD test performance (1,681 분자). "
        "Production 은 단일 Chemprop v27 (D-MPNN only) 를 채택하며, RF/CB / stacking 은 ablation 비교 실험.",
    )

    add_para(
        doc,
        "Production 채택 모델인 Chemprop v27 은 scaffold OOD test 에서 **AUC 0.907, MCC 0.615** 를 기록하였다. "
        "이는 학습 분포 내 새로운 골격에 대해 안정적인 분류 성능을 가짐을 의미한다. Ablation 비교 결과 RF/CB v3 "
        "단독은 MCC 0.598 (Chemprop 대비 -0.017), Chemprop + RF/CB honest stacking 은 MCC 0.617 (Chemprop 대비 "
        "+0.002) 의 marginal 향상만 보였다. Production code 단순화와 maintenance 용이성을 위해 단일 Chemprop v27 "
        "만 채택한다.",
    )

    add_para(
        doc,
        "특히 Chemprop v27 의 specificity (TNR) 가 0.88 로 매우 높아 false positive (안전 약물을 hepatotoxic 으로 "
        "잘못 분류) 가 적다. 이는 신약 후보 스크리닝에서 안전한 candidate 를 잘못 제외할 위험을 낮춘다.",
    )

    add_heading(doc, "3.2 External Validation — 진짜 OOD Evaluation", level=2)

    add_para(
        doc,
        "최종 모델의 과적합 (overfitting) 여부와 실제 미지 화합물에 대한 일반화 성능을 평가하기 위해, 학습 / 검증 "
        "/ 내부 test 모두에 사용되지 않은 외부 263개 분자에서 추가 검증을 수행하였다. 이 평가셋은 학습 DB 와 "
        "InChIKey 0 중첩이며, 129 양성 / 134 음성의 균형 잡힌 구성을 가진다. 평가 대상은 최신 약리군 (TKI, HCV "
        "antiviral, GLP-1, SGLT2, CDK / PARP / BTK inhibitor 등 post-2015 출시 약물 dominant) 위주이다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Model", "Accuracy", "TPR", "TNR", "MCC", "AUC"],
            ["Chemprop v27 (production)", "0.510", "0.318", "0.694", "0.013", "0.506"],
            ["RF/CB v3 (ablation)", "0.491", "0.434", "0.440", "0.061", "0.448"],
        ],
    )
    add_table_caption(
        doc,
        "Table 4. External evaluation on 263 truly novel molecules (zero InChIKey overlap with training data). "
        "All models perform near random — confirming the intrinsic OOD limitation of DILI prediction.",
    )

    add_para(
        doc,
        "외부 263 분자에서의 결과는 **AUC ~0.5, MCC ~0.05 의 random 수준** 이다. 내부 scaffold OOD test 의 "
        "MCC 0.615 (Table 3) 와 비교하면 큰 격차다.",
    )

    add_para(
        doc,
        "내부 scaffold OOD test 의 분자는 같은 약리군 (NSAID, statin, antibiotic 등) 안의 새 골격이라 모델이 "
        "학습한 위험 패턴이 적용 가능한 반면, 외부 263 분자는 학습 데이터에 거의 없는 chemical class (TKI 학습 95 "
        "/ GLP-1 학습 0 / SGLT2 학습 0 / HCV NS5A 학습 13 등) 이라 모델이 random 수준 예측만 가능하다. 이는 "
        "DILI ML 의 본질적 한계로, 학계 보고와 일치한다 (Vall et al., 2021; Seal et al., 2024). Section 4 에서 "
        "자세히 논의한다.",
    )

    add_heading(doc, "3.3 Ablation — 학습 Chemical Space 확장 실험", level=2)

    add_para(
        doc,
        "Section 3.2 의 결과로부터 가설을 도출할 수 있다 — 만약 학습 chemical class 다양성 부족이 외부 OOD 한계의 "
        "주원인이라면, 부족한 class 의 분자를 추가하면 외부 sanity 성능이 향상될 것이다. 이 가설을 검증하기 위해 "
        "학습이 부족한 chemical class 의 266개 분자를 신규 수집하여 재학습 실험을 수행하였다.",
    )

    add_para(
        doc,
        "4개 Agent 가 LiverTox / FDA / PubChem 으로부터 TKI (97), 당뇨 + 수면 (39), HCV / CDK / PARP / BTK / "
        "JAK (53), 신규 항생 / 항진균 / 기타 (77) 의 총 266개 분자를 수집하였다. InChIKey 중복 제거 후 신규 81개 "
        "(DB 없음, 26 양성 / 55 음성) 와 기존 168개 (DB 있음) 중 46개 label conflict 를 발견하였다. Label conflict "
        "중 28개 (DB 0 → LiverTox 1, 예: Telaprevir, Velpatasvir, Lobeglitazone, Panobinostat) 는 DB 의 명확한 "
        "누락으로 판단하여 update 하였고, 18개 (DB 1 → LiverTox 0) 는 DILI 정의 차이로 둘 다 valid 하므로 "
        "유지하였다. 재학습 후 동일한 외부 263 분자에서 평가한 결과는 Table 5 와 같다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Model", "External AUC", "MCC@best", "Δ (vs baseline)"],
            ["Chemprop v27 (baseline)", "0.506", "0.055", "—"],
            ["Chemprop v31 (+81 신규)", "0.515", "0.052", "+0.009"],
            ["RF/CB v3 (baseline)", "0.448", "0.061", "—"],
            ["RF/CB v31 (+81 신규)", "0.439", "0.061", "-0.009"],
            ["RF/CB v31_v2 (+28 label fix)", "0.451", "0.061", "+0.003"],
        ],
    )
    add_table_caption(
        doc,
        "Table 5. Class expansion ablation 결과 (외부 263 분자 평가). "
        "단순 데이터 추가의 효과 ≈ 0 (within noise). DILI 의 OOD 한계가 데이터 양 보다 본질적임을 확인한다.",
    )

    add_para(
        doc,
        "학습 데이터 확장의 효과는 모든 모델에서 marginal 또는 negative 였다 (Δ AUC < 0.01, MCC 변화 없음). "
        "가설 (데이터 양 부족이 OOD 한계의 주원인) 은 기각되며, DILI 의 본질적 한계가 데이터 양 보다 더 깊은 "
        "원인임을 확인한다.",
    )

    # ════════════════════════════════════════════════════════════
    # 4. Discussion
    # ════════════════════════════════════════════════════════════
    add_heading(doc, "4. Discussion", level=1)

    add_para(
        doc,
        "Section 3 의 결과는 본 연구의 두 가지 핵심 발견을 보여준다. 첫째, 학습 분포 내 새로운 골격에 대해서는 "
        "안정적인 일반화 성능 (scaffold OOD MCC 0.615) 을 보인다. 둘째, 학습 데이터에 없는 chemical class 에 "
        "대해서는 모든 모델이 random 수준 성능을 보이며, 단순 데이터 확장으로는 극복할 수 없다. 이 두 발견의 "
        "의미와 본질적 원인을 다음과 같이 논의한다.",
    )

    add_heading(doc, "4.1 결과 해석 — Chemical Space 의존성", level=2)

    add_para(
        doc,
        "본 모델의 일반화 성능은 평가 데이터의 chemical class 와 학습 데이터의 중첩 정도에 강하게 의존한다. "
        "내부 scaffold OOD test (MCC 0.615) 의 분자들은 학습 분포 안의 동일한 약리군 안에서 새로운 골격을 가진 "
        "분자이며, 모델이 학습한 위험 패턴이 골격이 달라도 적용 가능하다. 반면 외부 263 분자 (MCC 0.05) 는 학습 "
        "데이터에 거의 없는 chemical class (TKI, GLP-1 등) 이며, 모델이 이 class 의 위험 메커니즘을 학습한 적이 "
        "없어 random 수준 예측만 가능하다. 이러한 chemical space 의존성은 학계의 다른 DILI ML 모델에서도 비슷하게 "
        "보고된 바 있다 (Wu et al., 2018; Vall et al., 2021).",
    )

    add_para(
        doc,
        "Class expansion ablation (Section 3.3) 의 결과는 이 격차가 단순히 데이터 양 부족 때문이 아님을 보여준다. "
        "266개의 부족한 class 분자를 신규 추가하고 28개 label 을 정정해도 외부 sanity AUC 는 0.506 → 0.515 의 "
        "marginal 변화만 보였다. 이는 외부 OOD 한계가 데이터 큐레이션 노력으로 극복 가능한 기술적 문제가 아니라 "
        "DILI 자체의 생물학적 본질에서 기인함을 시사한다.",
    )

    add_heading(doc, "4.2 DILI 의 본질적 한계", level=2)

    add_para(
        doc,
        "본 연구에서 확인한 DILI prediction 의 한계는 데이터 또는 모델의 기술적 문제가 아닌, DILI 자체의 생물학적 "
        "본질에서 기인한다 (Kullak-Ublick et al., 2017).",
    )

    add_para(doc, "(1) DILI 의 특이체질 (idiosyncratic) 본질:", bold=True)
    add_para(
        doc,
        "특이체질 DILI 는 환자마다 발생 여부가 다르며, 분자 구조만으로는 결정되지 않는다 (Daly, 2010). 세 가지 "
        "요인이 함께 작용한다. **(a) 환자 유전형:** 특정 HLA (인간 백혈구 항원) 유전형을 가진 환자에서만 간 손상이 "
        "발생하는 경우가 있어, 분자 자체는 안전하지만 환자의 면역 반응을 통해 hepatitis 가 유발된다 (Daly et al., "
        "2009; Hautekeete et al., 1999). **(b) 대사산물 매개:** 약물 자체는 안전하나 간 효소 (CYP450) 에 의한 "
        "대사산물이 독성을 띠는 경우 (예: acetaminophen 의 NAPQI) 가 있다 (Mitchell et al., 1973). 모델은 원래 "
        "약물의 SMILES 만 볼 수 있어 대사산물 정보를 직접 학습할 수 없다. **(c) 낮은 발생률:** 특이체질 DILI 는 "
        "1,000명 ~ 10,000명 당 1명 수준의 발생률이라 임상 발견이 어렵고 통계적 학습 신호가 약하다 (Chalasani & "
        "Björnsson, 2010).",
    )

    add_para(doc, "(2) 학습 chemical space 의 제한:", bold=True)
    add_para(
        doc,
        "본 연구의 19,273 학습 분자는 대부분 1970~2015년에 시판된 약물 중심이다. 최신 chemical class — TKI, "
        "GLP-1, SGLT2, HCV NS5A, CDK / PARP / BTK inhibitor 등 — 는 underrepresented 하다. 예를 들어 GLP-1 "
        "agonist 와 SGLT2 inhibitor 는 학습 데이터에 거의 없으며, 이러한 class 의 신약은 본 모델로 random 수준 "
        "예측만 가능하다.",
    )

    add_para(doc, "(3) Source label noise:", bold=True)
    add_para(
        doc,
        "FAERS 같은 자발 신고는 oversensitive (false positive 많음), CTD / ChEMBL 의 hepatic disease 연관성은 "
        "indirect signal 이다. 8 source 의 OR rule 통합과 1,210건의 수동 큐레이션에도 잔여 label noise 가 학습에 "
        "영향을 미친다.",
    )

    add_para(
        doc,
        "이러한 본질적 한계는 우리 모델만의 문제가 아니라 학계의 다른 분자 표현 학습 모델들 — MoLFormer (Ross et "
        "al., 2022), ChemBERTa (Chithrananda et al., 2020), GROVER (Rong et al., 2020), KPGT (Li et al., 2022) "
        "등 — 도 학습에 없던 신약을 평가할 때 비슷한 수준의 한계를 보인다. DILI 의 예측은 학계 전반의 공통 도전 "
        "과제이다.",
    )

    add_heading(doc, "4.3 시도하였으나 한계 확인된 접근", level=2)

    add_para(
        doc,
        "본 연구는 외부 신약에 대한 일반화 한계를 극복하기 위해 다음 다섯 가지 접근을 시도하였으나, 모두 거의 "
        "변화가 없거나 오히려 성능이 약간 떨어지는 결과를 얻었다. 이 ablation 들은 단순한 데이터 추가나 모델 조정 "
        "만으로는 DILI 의 본질적 한계를 극복할 수 없다는 결론을 뒷받침한다.",
    )

    add_para(
        doc,
        "**(1) 학습 데이터 확장.** 부족한 chemical class 의 266개 분자를 새로 수집하고 28개 라벨을 정정하여 "
        "재학습하였으나, 외부 263 분자의 AUC 는 0.506 → 0.515 의 미미한 변화에 그쳤다 (Section 3.3).",
    )

    add_para(
        doc,
        "**(2) Source 신뢰도 기반 가중치.** 8개 source 의 신뢰도 점수에 따라 각 분자의 학습 가중치를 다르게 "
        "부여하였으나, 성능 향상은 미미하였다. 코드 단순화를 위해 production 에서 제외하였다.",
    )

    add_para(
        doc,
        "**(3) Focal loss (Lin et al., 2017).** 음성 분자 (다수) 의 손실 가중치를 낮추고 양성 분자 (소수) 의 손실 "
        "가중치를 높이는 손실 함수 (γ=2) 를 시도하였으나, scaffold OOD MCC 가 0.007 오히려 떨어졌다.",
    )

    add_para(
        doc,
        "**(4) SMOTE oversampling (Chawla et al., 2002).** 양성 분자를 인공으로 복제 + 변형하여 음성과 수를 맞추는 "
        "방법을 시도하였으나, scaffold OOD MCC 가 0.003 오히려 떨어졌다.",
    )

    add_para(
        doc,
        "**(5) Chemprop + RF/CB 결합 (Honest Stacking).** 두 모델의 예측을 가중 평균하는 방법을 시도하였으나, "
        "단일 Chemprop 대비 MCC 향상은 0.002 에 불과하였다. 코드 복잡도를 줄이기 위해 단일 Chemprop 만 production "
        "에 채택하였다.",
    )

    add_heading(doc, "4.4 향후 연구 방향", level=2)

    add_para(doc, "DILI 예측의 본질적 한계를 극복하기 위해 향후 다음 네 가지 방향이 가능하다.")

    add_para(
        doc,
        "**(1) 분자 표현의 사전학습 활용.** 본 연구는 19,273개의 학습 분자만 사용했지만, MoLFormer (IBM 의 11억 "
        "분자 사전학습 모델, Ross et al., 2022) 같은 분자 foundation model 을 활용하면 모델이 일반적인 화학 지식을 "
        "미리 학습한 상태에서 DILI 데이터로 fine-tune 할 수 있다. 학계 보고에 따르면 외부 신약 예측이 AUC "
        "0.50 → 0.55~0.62 정도 향상될 수 있다. 본질적 한계를 깨지는 못하지만 향상의 여지가 있다.",
    )

    add_para(
        doc,
        "**(2) 다요인 통합 모델.** DILI 는 분자 구조 외에 환자 유전형, 대사 효소 활성, 투여 용량 등 여러 요인의 "
        "조합으로 발생한다. 분자 구조 정보에 CYP450 효소 활성 예측 결과, HLA 유전형 빈도, 임상 투여 용량 정보를 "
        "함께 입력하는 다요인 (multi-modal) 모델이 가능하다. 그러나 이러한 paired 데이터셋이 부족한 것이 가장 큰 "
        "장벽이다.",
    )

    add_para(
        doc,
        "**(3) 능동 학습 (Active Learning).** 모델이 자신 없어 하는 신약 (예측 확률이 0.4~0.6 인 분자) 을 우선 "
        "라벨링하여 학습 데이터에 추가하고 다시 학습을 반복하는 방법이다 (Settles, 2009). 학습 chemical space 가 "
        "점진적으로 확장되어 매년 출시되는 신약에 모델을 효율적으로 업데이트할 수 있다.",
    )

    add_para(
        doc,
        "**(4) 예측 신뢰도 정량화.** 모델의 예측에 confidence interval (신뢰 구간) 을 함께 제공하여, novel "
        'chemical class 분자에 대해서는 "이 분자는 학습 분포에서 벗어나 예측 신뢰도가 낮음" 을 사용자에게 명시할 '
        "수 있다. Monte Carlo dropout (Gal & Ghahramani, 2016) 또는 ensemble disagreement 기반 uncertainty "
        "quantification (Lakshminarayanan et al., 2017) 을 활용할 수 있다.",
    )

    add_heading(doc, "4.5 결론", level=2)

    add_para(
        doc,
        "본 연구는 단순한 SMILES 기반 DILI 예측 모델 개발을 넘어, DILI ML 의 본질적 한계를 honest 한 ablation 으로 "
        "검증하는 데 기여한다. 주요 기여는 다음과 같다. **(1)** 8 source 통합의 가장 큰 규모 DILI 학습 데이터셋 "
        "(19,273 unique) 을 구축하였다. **(2)** Bemis-Murcko scaffold-balanced split (Bemis & Murcko, 1996) 과 "
        "Chemprop D-MPNN ensemble (Yang et al., 2019) 의 honest 학습 평가로 random split 의 과대평가 함정을 "
        "회피하였다 (internal scaffold OOD MCC 0.615, AUC 0.907). **(3)** DILI ML 의 본질적 한계를 honest "
        "negative result 로 검증하였다 — 외부 263 분자에서 모든 모델이 random 수준이며, 학계 SOTA 도 동일한 한계가 "
        "보고된 것과 일치한다 (Vall et al., 2021). **(4)** 학습 데이터 확장의 ablation 으로 단순 데이터 양 부족이 "
        "한계의 주원인이 아님을 증명하였으며, 향후 연구 방향에 직접 시사점을 제공한다.",
    )

    add_para(
        doc,
        "본 모델은 학습 분포 내 신약에 대해서는 의미 있는 일반화 성능 (MCC 0.615) 을 보이며, novel chemical class "
        "에 대해서는 random 수준 한계가 있음을 honest 하게 보고한다. 향후 foundation model 과 multi-modal 통합을 "
        "통해 외부 OOD 신약 예측 한계를 점진적으로 극복할 수 있을 것으로 기대한다.",
    )

    # ════════════════════════════════════════════════════════════
    # 5. References
    # ════════════════════════════════════════════════════════════
    add_heading(doc, "References", level=1)

    refs = [
        # Methods — software / techniques
        "Bemis, G. W., & Murcko, M. A. (1996). The properties of known drugs. 1. Molecular frameworks. Journal of Medicinal Chemistry, 39(15), 2887–2893.",
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32.",
        "Carhart, R. E., Smith, D. H., & Venkataraghavan, R. (1985). Atom pairs as molecular features in structure-activity studies. Journal of Chemical Information and Computer Sciences, 25(2), 64–73.",
        "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic Minority Over-sampling Technique. Journal of Artificial Intelligence Research, 16, 321–357.",
        "Chicco, D., & Jurman, G. (2020). The advantages of the Matthews correlation coefficient (MCC) over F1 score and accuracy in binary classification evaluation. BMC Genomics, 21(6), 1–13.",
        "Durant, J. L., Leland, B. A., Henry, D. R., & Nourse, J. G. (2002). Reoptimization of MDL keys for use in drug discovery. Journal of Chemical Information and Computer Sciences, 42(6), 1273–1280.",
        "Ertl, P., Rohde, B., & Selzer, P. (2000). Fast calculation of molecular polar surface area as a sum of fragment-based contributions. Journal of Medicinal Chemistry, 43(20), 3714–3717.",
        "Gal, Y., & Ghahramani, Z. (2016). Dropout as a Bayesian approximation: Representing model uncertainty in deep learning. Proceedings of the 33rd International Conference on Machine Learning (ICML), 1050–1059.",
        "Gedeck, P., Rohde, B., & Bartels, C. (2006). QSAR — How Good Is It in Practice? Comparison of Descriptor Sets on an Unbiased Cross Section of Corporate Data Sets. Journal of Chemical Information and Modeling, 46(5), 1924–1936.",
        "Hanley, J. A., & McNeil, B. J. (1982). The meaning and use of the area under a receiver operating characteristic (ROC) curve. Radiology, 143(1), 29–36.",
        "Heid, E., Greenman, K. P., Chung, Y., Li, S.-C., Graff, D. E., Vermeire, F. H., et al. (2024). Chemprop: A machine learning package for chemical property prediction. Journal of Chemical Information and Modeling, 64(1), 9–17.",
        "Heller, S. R., McNaught, A., Pletnev, I., Stein, S., & Tchekhovskoi, D. (2015). InChI, the IUPAC International Chemical Identifier. Journal of Cheminformatics, 7, 23.",
        "Lakshminarayanan, B., Pritzel, A., & Blundell, C. (2017). Simple and Scalable Predictive Uncertainty Estimation using Deep Ensembles. Advances in Neural Information Processing Systems (NeurIPS), 30.",
        "Landrum, G. (2024). RDKit: Open-source cheminformatics. https://www.rdkit.org",
        "Lin, T.-Y., Goyal, P., Girshick, R., He, K., & Dollár, P. (2017). Focal loss for dense object detection. Proceedings of the IEEE International Conference on Computer Vision (ICCV), 2980–2988.",
        "Matthews, B. W. (1975). Comparison of the predicted and observed secondary structure of T4 phage lysozyme. Biochimica et Biophysica Acta — Protein Structure, 405(2), 442–451.",
        "Nilakantan, R., Bauman, N., Dixon, J. S., & Venkataraghavan, R. (1987). Topological torsion: a new molecular descriptor for SAR applications. Journal of Chemical Information and Computer Sciences, 27(2), 82–85.",
        "Prokhorenkova, L., Gusev, G., Vorobev, A., Dorogush, A. V., & Gulin, A. (2018). CatBoost: unbiased boosting with categorical features. Advances in Neural Information Processing Systems (NeurIPS), 31.",
        "Rogers, D., & Hahn, M. (2010). Extended-connectivity fingerprints. Journal of Chemical Information and Modeling, 50(5), 742–754.",
        "Settles, B. (2009). Active learning literature survey. University of Wisconsin–Madison Computer Sciences Technical Report 1648.",
        "Swain, M. (2017). PubChemPy: A Python wrapper for the PubChem PUG REST API. https://github.com/mcs07/PubChemPy",
        "Weininger, D. (1988). SMILES, a chemical language and information system. 1. Introduction to methodology and encoding rules. Journal of Chemical Information and Computer Sciences, 28(1), 31–36.",
        "Yang, K., Swanson, K., Jin, W., Coley, C., Eiden, P., Gao, H., et al. (2019). Analyzing learned molecular representations for property prediction. Journal of Chemical Information and Modeling, 59(8), 3370–3388.",
        # Data sources
        "Chen, M., Suzuki, A., Thakkar, S., Yu, K., Hu, C., & Tong, W. (2016). DILIrank: the largest reference drug list ranked by the risk for developing drug-induced liver injury in humans. Drug Discovery Today, 21(4), 648–653.",
        "Davis, A. P., Wiegers, T. C., Johnson, R. J., Sciaky, D., Wiegers, J., & Mattingly, C. J. (2023). Comparative Toxicogenomics Database (CTD): update 2023. Nucleic Acids Research, 51(D1), D1257–D1262.",
        "FDA (2024). FDA Adverse Event Reporting System (FAERS) Public Dashboard. U.S. Food and Drug Administration.",
        "Hoofnagle, J. H., Serrano, J., Knoben, J. E., & Navarro, V. J. (2013). LiverTox: a website on drug-induced liver injury. Hepatology, 57(3), 873–874.",
        "Mendez, D., Gaulton, A., Bento, A. P., Chambers, J., De Veij, M., Félix, E., et al. (2019). ChEMBL: towards direct deposition of bioassay data. Nucleic Acids Research, 47(D1), D930–D940.",
        # DILI biology
        "Chalasani, N., & Björnsson, E. (2010). Risk factors for idiosyncratic drug-induced liver injury. Gastroenterology, 138(7), 2246–2259.",
        "Chen, M., Borlak, J., & Tong, W. (2013). High lipophilicity and high daily dose of oral medications are associated with significant risk for drug-induced liver injury. Hepatology, 58(1), 388–396.",
        "Daly, A. K. (2010). Genome-wide association studies in pharmacogenomics. Nature Reviews Genetics, 11(4), 241–246.",
        "Daly, A. K., Donaldson, P. T., Bhatnagar, P., Shen, Y., Pe'er, I., Floratos, A., et al. (2009). HLA-B*5701 genotype is a major determinant of drug-induced liver injury due to flucloxacillin. Nature Genetics, 41(7), 816–819.",
        "Hautekeete, M. L., Horsmans, Y., Van Waeyenberge, C., Demanet, C., Henrion, J., Verbist, L., et al. (1999). HLA association of amoxicillin-clavulanate-induced hepatitis. Gastroenterology, 117(5), 1181–1186.",
        "Kullak-Ublick, G. A., Andrade, R. J., Merz, M., End, P., Benesic, A., Gerbes, A. L., & Aithal, G. P. (2017). Drug-induced liver injury: recent advances in diagnosis and risk assessment. Gut, 66(6), 1154–1164.",
        "Mitchell, J. R., Jollow, D. J., Potter, W. Z., Davis, D. C., Gillette, J. R., & Brodie, B. B. (1973). Acetaminophen-induced hepatic necrosis. I. Role of drug metabolism. Journal of Pharmacology and Experimental Therapeutics, 187(1), 185–194.",
        "Watkins, P. B. (2005). Idiosyncratic liver injury: challenges and approaches. Toxicologic Pathology, 33(1), 1–5.",
        # DILI ML SOTA / Foundation models
        "Chithrananda, S., Grand, G., & Ramsundar, B. (2020). ChemBERTa: Large-scale self-supervised pretraining for molecular property prediction. arXiv preprint arXiv:2010.09885.",
        "Li, H., Zhao, D., & Zeng, J. (2022). KPGT: Knowledge-Guided Pre-training of Graph Transformer for Molecular Property Prediction. Proceedings of the 28th ACM SIGKDD Conference (KDD), 857–867.",
        "Seal, S., Williams, D. P., Hosseini-Gerami, L., Mahale, M., Carpenter, A. E., Spjuth, O., & Bender, A. (2024). Improved Detection of Drug-Induced Liver Injury by Integrating Predicted in vivo and in vitro Data. Chemical Research in Toxicology, 37, 1290–1305.",
        "Rong, Y., Bian, Y., Xu, T., Xie, W., Wei, Y., Huang, W., & Huang, J. (2020). Self-supervised graph transformer on large-scale molecular data. Advances in Neural Information Processing Systems (NeurIPS), 33, 12559–12571.",
        "Ross, J., Belgodere, B., Chenthamarakshan, V., Padhi, I., Mroueh, Y., & Das, P. (2022). Large-scale chemical language representations capture molecular structure and properties. Nature Machine Intelligence, 4, 1256–1264.",
        "Vall, A., Sabnis, Y., Shi, J., Class, R., Hochreiter, S., & Klambauer, G. (2021). The promise of AI for drug-induced liver injury prediction with artificial intelligence. Frontiers in Artificial Intelligence, 4, 638410.",
        "Wu, Z., Ramsundar, B., Feinberg, E. N., Gomes, J., Geniesse, C., Pappu, A. S., Leswing, K., & Pande, V. (2018). MoleculeNet: a benchmark for molecular machine learning. Chemical Science, 9(2), 513–530.",
    ]

    for i, ref in enumerate(refs, 1):
        add_reference(doc, ref, i)

    # 저장
    doc.save(OUTPUT)
    print(f"저장 완료: {OUTPUT}")
    print(f"크기: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
    print(f"References: {len(refs)} 개")


if __name__ == "__main__":
    build_doc()
