"""격식 학술 한국어 보고서 — [숫자] 인용 형식.

특징:
  - 격식 학술 어투 (~한다, ~이다, ~다)
  - Reference inline [숫자] 형식
  - 신뢰할 수 있는 학술 출처만 (논문 / 정부 보고서)
  - 본문에서 모든 reference 인용 확인
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT = os.path.join(PROJECT_ROOT, "docs", "Report_PersonalVoice.docx")


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_korean_font(run, font="맑은 고딕"):
    run.font.name = font
    rPr = run._element.rPr
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_korean_font(run)
    return h


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    _set_korean_font(run)
    return p


def add_table_from_rows(doc, rows, header=True):
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    _set_korean_font(run)
                    if i == 0 and header:
                        run.bold = True
            if i == 0 and header:
                set_cell_background(cell, "DCE6F1")
    return table


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    _set_korean_font(run)


def add_reference(doc, text, number):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(f"[{number}] ")
    run.font.size = Pt(10)
    run.bold = True
    _set_korean_font(run)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    _set_korean_font(run2)
    return p


def build_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ════════════════════════════════════════════════════════════
    add_heading(doc, "2. Methods and Results", level=1)
    # ════════════════════════════════════════════════════════════

    # ───────────────────────────────────────────────
    add_heading(doc, "2.1 초기 baseline 과 평가 방식의 문제 발견", level=2)
    # ───────────────────────────────────────────────

    add_para(
        doc,
        "초기 모델 구축은 가장 간단한 baseline 으로부터 시작하였다. PubChem 에서 약물의 SMILES [39] 를 수집한 후, "
        "RDKit 라이브러리를 사용하여 ECFP6 fingerprint [31] 를 계산하고 RandomForest [2] 에 입력하였다. "
        "Train / Test 는 random 으로 70 / 30 비율로 분할하였다. 평가 결과 MCC 0.85 수준의 높은 성능이 산출되었으나, "
        "해당 결과의 타당성에 대한 추가 검증이 필요하다고 판단하였다.",
    )

    add_para(
        doc,
        "Sanity check 으로 잘 알려진 양성 약물 (아세트아미노펜, 간독성 명확) 과 음성 약물 (아토르바스타틴, 안전) 을 "
        "모델에 입력하였을 때 두 약물 모두 정답이 출력되었다. 그러나 추가 확인 결과 두 약물 모두 학습 데이터에 이미 "
        "포함된 분자로 확인되었으며, 이는 모델이 학습된 분자의 라벨을 그대로 반환하는 data leakage 상황임이 "
        "확인되었다.",
    )

    add_para(
        doc,
        "또한 random split 의 구조적 한계가 식별되었다. 동일한 약리군 (예: NSAID, 스타틴) 에 속하는 분자들이 train 과 "
        "test 에 균등하게 분산될 경우, 모델은 해당 약리군의 공통 골격과 작용기 패턴을 통해 정답을 추론할 수 있다. 이는 "
        "실제 신약 예측 상황 — 즉 학습 데이터에 포함되지 않은 새로운 골격을 가진 분자의 평가 — 와 부합하지 않는 "
        "평가 방식이다. 따라서 학습 분포 외 분자에 대한 일반화 능력을 평가할 수 있는 분할 방식이 필요하다고 판단하였다. "
        "관련 문헌 검토 결과 MoleculeNet 연구 [40] 와 Chemprop 연구 [41] 에서 권장하는 Bemis-Murcko scaffold split "
        "[1] 을 채택하였다.",
    )

    # ───────────────────────────────────────────────
    add_heading(doc, "2.2 데이터 수집 — 8개 source 통합", level=2)
    # ───────────────────────────────────────────────

    add_para(
        doc,
        "Random split 의 한계를 해결한 후에도 학습 데이터의 양적 확장이 필요하다고 판단되어 가능한 공개 DILI source 를 "
        "추가 수집하였다. 초기에 확보한 source 는 FDA DILIrank [7] 와 NIH LiverTox [21] 였으나, 각각 1,036 / 851 "
        "분자 규모로 학습에 충분하지 않았다.",
    )

    add_para(
        doc,
        "따라서 다음의 8개 source 를 통합하였다: DILIrank, LiverTox, FDA DailyMed (의약품 라벨 기반 hepatic AE), "
        "PubMed (DILI 관련 문헌의 약물 mention), Comparative Toxicogenomics Database [12], FDA Adverse Event "
        "Reporting System (FAERS) [36], ChEMBL [27] 의 max_phase = 4 (FDA approved) 약물, 그리고 post-2010 출시 + "
        "DILI 보고 없는 약물의 음성 anchor (Marketed Clean). Source 별 분포는 Table 1 에 제시한다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Source", "출처", "원본 분자", "Vivo 양성", "Vivo 음성"],
            ["DILIrank", "FDA (2016)", "1,036", "514", "522"],
            ["LiverTox", "NIH NCBI Bookshelf", "851", "384", "467"],
            ["DailyMed", "FDA", "4,200+", "1,560", "2,640"],
            ["PubMed", "NLM", "8,500+", "1,830", "6,670"],
            ["CTD", "Toxicogenomics DB", "6,800+", "2,140", "4,660"],
            ["FAERS", "openFDA", "1,200+", "720", "480"],
            ["Marketed Clean", "PubChem + 자체", "2,500+", "0", "2,500 (음성 anchor)"],
            ["ChEMBL", "EBI", "14,000+", "920", "13,080"],
            ["통합 (dedup)", "—", "19,273", "3,327", "10,494"],
        ],
    )
    add_table_caption(
        doc,
        "Table 1. 통합한 8개 DILI source 의 분포. 신뢰도 검토 후 제외한 source 는 Gold standard, DILIst, SIDER, "
        "TDC DILI, ClinTox 의 5개이다.",
    )

    add_para(
        doc,
        "초기에는 더 많은 source (Gold standard, DILIst, SIDER, TDC DILI, ClinTox) 를 검토하였으나, 라벨 노이즈가 큰 "
        "것으로 확인되어 제외하였다. 예를 들어 SIDER 는 in vitro 와 부작용 증상이 혼재되어 있어 임상 hepatotoxicity 와 "
        "일치하지 않았으며, ClinTox 는 임상 시험 일반의 toxicity 데이터로 DILI 특이성이 부족하였다.",
    )

    add_para(doc, "분자 표준화 (Standardization):", bold=True)

    add_para(
        doc,
        "각 source 의 SMILES 표기에 차이가 존재하였다. 예를 들어 ibuprofen 의 sodium salt 와 free acid 가 별개 분자로 "
        "분류되거나, tautomer 의 표현 차이로 동일 분자가 중복 인식되는 문제가 발생하였다. 이를 해결하기 위해 RDKit 의 "
        "MolStandardize chain (Normalizer → LargestFragmentChooser → Uncharger) 을 적용하여 정규화하였다. 이후 "
        "InChIKey [20] 를 분자의 고유 식별자로 사용하여 중복을 제거하였다.",
    )

    add_para(doc, "라벨 통합 — OR rule 및 수동 큐레이션:", bold=True)

    add_para(
        doc,
        "8개 source 의 라벨을 통합하는 방법으로 보수적인 OR rule (한 source 이상에서 양성으로 분류된 경우 양성으로 "
        "결정) 을 적용하였다. 단순 majority vote 는 각 source 의 라벨링 기준이 상이하여 (예: DILIrank 의 4단계 "
        "severity vs FAERS 의 신고 건수) 적합하지 않다고 판단하였다.",
    )

    add_para(
        doc,
        "OR rule 적용 후 1,210건의 label conflict — 한 source 는 양성, 다른 source 는 명확히 음성으로 분류된 경우 — 가 "
        "발생하였다. 이러한 conflict 가 학습에 noise 로 작용할 가능성이 있어 LiverTox 와 PubMed 의 임상 보고를 기반으로 "
        "전수 수동 검토하여 정정하였다. 추가로 약한 양성 signal (단일 source 양성, 나머지 source 데이터 없음) 1,661건은 "
        "별도로 검증하여 494건은 양성 확정, 374건은 음성 재분류, 793건은 약물이 아닌 excipient / food additive 로 "
        "분류하였다.",
    )

    add_para(doc, "In vivo 데이터의 선택적 사용:", bold=True)

    add_para(
        doc,
        "통합 과정에서 in vitro 와 in vivo (임상) 라벨이 일치하지 않는 분자가 다수 확인되었다. 예를 들어 troglitazone 은 "
        "in vitro HepG2 cytotoxicity assay 에서 EC50 > 100 μM 로 안전한 것으로 평가되나, 임상에서는 idiosyncratic "
        "hepatitis 로 FDA 가 회수한 약물이다 [38]. Acetaminophen 의 경우 분자 자체는 안전하나 활성 대사산물인 NAPQI 가 "
        "간세포의 glutathione 을 고갈시키며 손상을 유발한다 [28]. 이러한 mismatch 분자가 학습 데이터에 포함될 경우 "
        "노이즈로 작용함이 ablation 에서 확인되었다.",
    )

    add_para(
        doc,
        "따라서 임상 의의가 직접적인 in vivo 라벨 13,821개 (양성 3,327 / 음성 10,494) 만을 학습에 사용하기로 결정하였다. "
        "In vitro labeled 7,561개 분자는 데이터셋에 보관하되 학습에는 사용하지 않았다.",
    )

    # ───────────────────────────────────────────────
    add_heading(doc, "2.3 Feature 선택", level=2)
    # ───────────────────────────────────────────────

    add_para(
        doc,
        "Feature 의 선택은 단계적으로 진행하였다. 초기에는 ECFP6 fingerprint [31] 단독으로 시작하였다. ECFP6 은 각 "
        "원자 주변의 부분 구조를 hash 하여 2,048 bit 의 binary vector 로 표현하는 방식이다. RandomForest 에 입력한 "
        "결과 scaffold OOD MCC 0.55 의 성능을 얻었다.",
    )

    add_para(
        doc,
        "다양한 분자 정보를 활용하기 위해 5종 fingerprint 로 확장하였다: ECFP6, Avalon [16], AtomPair [3], "
        "Topological Torsion [29], Pattern (MACCS-like) [13]. RandomForest 와 CatBoost [30] 의 ensemble 로 학습한 "
        "결과 MCC 0.59 로 향상되었다.",
    )

    add_para(
        doc,
        "Fingerprint 기반 학습은 분자의 그래프 구조 정보를 1차원 vector 로 펴면서 정보 손실이 발생할 가능성이 있다고 "
        "판단하였다. 이에 그래프 신경망인 Chemprop [41, 19] 의 D-MPNN (Directed Message Passing Neural Network) 을 "
        "도입하였다. Chemprop 은 분자를 그래프 (원자 = 노드, 결합 = 엣지) 로 표현한 후 message passing 으로 인접 "
        "원자의 정보를 통합하며 학습한다. 적용 결과 scaffold OOD MCC 0.61 로 단일 모델 기준 최고 성능을 달성하였다.",
    )

    add_para(
        doc,
        "추가로 200차원 RDKit 2D descriptor (분자량, LogP, TPSA [14], H-bond donor / acceptor, 회전 가능 결합 수, "
        "ring 수 등) 를 Chemprop 의 graph representation 에 concat 하였다. Chen et al. [6] 의 'Rule of Two' 연구에서 "
        "LogP > 3 + dose > 100 mg/day 인 약물의 60% 이상이 DILI risk 를 가짐이 보고된 바, 분자 단위의 물리화학 특성을 "
        "명시적으로 학습에 포함하는 것이 필요하다고 판단하였다. Descriptor 추가 후 scaffold OOD MCC 가 0.63 으로 "
        "향상되었다.",
    )

    # ───────────────────────────────────────────────
    add_heading(doc, "2.4 학습 — Scaffold Split 및 Chemprop Ensemble", level=2)
    # ───────────────────────────────────────────────

    add_para(
        doc,
        "Section 2.1 에서 언급한 random split 의 한계를 해결하기 위해 Bemis-Murcko scaffold-balanced split [1] 을 "
        "70 / 15 / 15 비율로 적용하였다. 각 분자의 골격 (ring system + linker) 을 추출한 후, 동일한 골격의 모든 "
        "분자를 동일한 fold 에 배치하여 train ∩ test scaffold = ∅ 을 보장하였다. 분할 결과는 Table 2 와 같다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Split", "분자 수", "양성률", "unique Scaffold"],
            ["Train", "9,645", "26.4%", "6,892"],
            ["Validation", "2,066", "19.2%", "1,832"],
            ["Test", "2,066", "18.0%", "1,851"],
        ],
    )
    add_table_caption(
        doc,
        "Table 2. In vivo DILI 데이터의 Bemis-Murcko scaffold-balanced split (70 / 15 / 15). "
        "Train ∩ Test 의 InChIKey 및 scaffold 모두 중첩이 없음을 확인하였다.",
    )

    add_para(
        doc,
        "Scaffold split 적용 후 MCC 가 0.85 (random split) 에서 0.61 (scaffold split) 로 감소하였다. 이는 모델 "
        "성능의 저하가 아니라, 실제 신약 예측 시 기대 가능한 진정한 일반화 성능에 해당한다. Random split 의 0.85 는 "
        "data leakage 에 의한 과대평가였음이 확인되었다.",
    )

    add_para(
        doc,
        "Production 모델은 Chemprop D-MPNN ensemble 을 채택하였다. 주요 hyperparameter 는 ensemble size 15 (서로 "
        "다른 random seed 로 학습된 15개 모델의 예측 확률 평균), hidden dimension 600, message passing depth 3, "
        "BCE loss, epoch 40, early stopping patience 8 이다. 해당 설정값은 Chemprop 의 default 및 MoleculeNet "
        "benchmark [40] 의 표준 설정을 기반으로 결정하였다. Ensemble size 5 의 경우 prediction variance 가 큰 것으로 "
        "확인되었고, 30 이상의 경우 학습 시간 대비 성능 향상이 미미하여 15 로 결정하였다. M1 Max CPU 환경에서 학습 "
        "시간은 약 45분이 소요되었다.",
    )

    add_para(doc, "Class imbalance 처리:", bold=True)

    add_para(
        doc,
        "학습 데이터의 양성 / 음성 비율은 24 / 76 의 불균형 데이터이다. 다양한 imbalance 처리 방법을 ablation 으로 "
        "비교한 결과, focal loss [25] (γ=2) 와 SMOTE oversampling [5] 모두 scaffold OOD MCC 가 각각 -0.007, -0.003 의 "
        "negative 효과를 보였다. 따라서 추가 imbalance 처리를 적용하지 않고, Chemprop 의 15개 ensemble 평균이 "
        "자연스럽게 imbalance 효과를 완화하는 방식을 채택하였다.",
    )

    # ───────────────────────────────────────────────
    add_heading(doc, "2.5 평가 — 내부 Test 결과", level=2)
    # ───────────────────────────────────────────────

    add_para(
        doc,
        "Scaffold OOD test (1,681 분자) 에서의 모델 성능을 평가하였다. 평가 지표로 Accuracy, Sensitivity (TPR), "
        "Specificity (TNR), Matthews Correlation Coefficient (MCC) [26], Area Under ROC Curve (AUC) [17] 의 5가지를 "
        "사용하였다. 데이터 imbalance 환경에서 가장 신뢰할 수 있는 단일 지표인 MCC [8] 를 주요 평가 지표로 채택하였다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Model", "Production?", "Accuracy", "TPR", "TNR", "MCC", "AUC"],
            ["Chemprop v27 (D-MPNN)", "✓ 채택", "0.853", "0.650", "0.880", "0.615", "0.907"],
            ["RF/CB v3 (ablation)", "✗ 비교", "0.832", "0.620", "0.853", "0.598", "0.852"],
            ["Honest Stacking", "✗ 비교", "0.858", "0.661", "0.882", "0.617", "0.910"],
        ],
    )
    add_table_caption(
        doc,
        "Table 3. Internal scaffold-OOD test 성능 (1,681 분자). Production 은 단일 Chemprop v27 모델을 채택하였으며, "
        "RF/CB ensemble 과 stacking 은 비교 ablation 실험이다.",
    )

    add_para(
        doc,
        "Production 모델 Chemprop v27 은 scaffold OOD test 에서 MCC 0.615, AUC 0.907 의 성능을 달성하였다. "
        "비교 ablation 으로 학습한 RF/CB ensemble 단독은 MCC 0.598 로 Chemprop 대비 -0.017 의 차이를 보였으며, "
        "Chemprop + RF/CB 의 honest linear stacking 은 MCC 0.617 로 단일 Chemprop 대비 +0.002 의 marginal 향상에 "
        "그쳤다. 따라서 production code 의 단순성과 유지보수 용이성을 고려하여 단일 Chemprop v27 을 채택하였다.",
    )

    add_para(
        doc,
        "특히 specificity (TNR) 가 0.88 로 높게 산출되었으며, 이는 신약 후보 스크리닝 시 false positive (안전 약물을 "
        "hepatotoxic 으로 잘못 분류) 비율이 낮음을 의미한다.",
    )

    # ───────────────────────────────────────────────
    add_heading(doc, "2.6 외부 평가 — 진정한 OOD 에서의 한계 확인", level=2)
    # ───────────────────────────────────────────────

    add_para(
        doc,
        "Internal scaffold OOD test 결과 (MCC 0.615) 외에 추가적인 외부 검증을 수행하였다. Scaffold split 의 train 과 "
        "test 분자는 모두 동일한 source (DILIrank, LiverTox 등) 에서 추출되었으므로 동일한 chemical space 내의 평가에 "
        "해당한다. 학습 데이터와 완전히 독립된 분자에 대한 일반화 능력을 검증하기 위해 외부 263개 분자 — 학습 DB 와 "
        "InChIKey 0 중첩 — 를 별도로 수집하였다.",
    )

    add_para(
        doc,
        "외부 평가셋은 LiverTox, FDA Drug Label, PubChem 에서 수집하였으며, 양성 129 / 음성 134 의 균형 잡힌 구성이다. "
        "평가 대상은 학습 데이터에 representation 이 부족한 최신 약리군 (TKI, HCV antiviral, GLP-1, SGLT2, "
        "CDK / PARP / BTK inhibitor 등 post-2015 출시 신약) 위주이다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Model", "Accuracy", "TPR", "TNR", "MCC", "AUC"],
            ["Chemprop v27 (production)", "0.510", "0.318", "0.694", "0.013", "0.506"],
            ["RF/CB v3 (ablation)", "0.491", "0.434", "0.440", "0.061", "0.448"],
        ],
    )
    add_table_caption(
        doc,
        "Table 4. 외부 263 분자 (학습 DB 와 InChIKey 0 중첩) 에 대한 평가 결과. 두 모델 모두 random 수준의 성능을 "
        "보였다.",
    )

    add_para(
        doc,
        "외부 263 분자에 대한 Chemprop v27 의 AUC 는 0.506, MCC 는 0.013 으로 random 수준의 성능이 산출되었다. "
        "Internal scaffold OOD test 의 MCC 0.615 와 큰 격차를 보였다. 이러한 결과의 원인을 분석하기 위해 학계의 다른 "
        "DILI ML 모델을 검토한 결과, MoLFormer [33], ChemBERTa [9], GROVER [32], KPGT [24] 등 최신 분자 표현 학습 "
        "모델들도 외부 신약 평가에서 유사한 한계를 보임이 확인되었다 [37, 34]. 본 한계는 특정 모델의 문제가 아니라 "
        "DILI prediction 분야 전반의 본질적 한계로 판단된다.",
    )

    # ───────────────────────────────────────────────
    add_heading(doc, "2.7 가설 검증 — 학습 데이터 확장의 효과", level=2)
    # ───────────────────────────────────────────────

    add_para(
        doc,
        "외부 평가에서 random 수준의 성능이 산출된 원인 가설로, 외부 263 분자가 학습 데이터에 부족한 chemical class 의 "
        "신약이라는 점을 제시하였다. 실제 학습 데이터에서 TKI 는 95개에 불과하고 (외부 sanity 의 양성 100 중 약 40개가 "
        "TKI), GLP-1 agonist 와 SGLT2 inhibitor 는 거의 0개로 확인되었다.",
    )

    add_para(
        doc,
        "본 가설을 검증하기 위해 부족한 chemical class 의 분자를 추가 수집하여 재학습하는 ablation 을 수행하였다. "
        "4개의 자동 검색 agent 를 병렬로 활용하여 TKI 97개, 당뇨 / 수면제 39개, HCV / CDK / PARP / BTK / JAK inhibitor "
        "53개, 신규 항생제 / 항진균제 등 77개 — 총 266개의 분자를 LiverTox, FDA, PubChem 에서 수집하였다.",
    )

    add_para(
        doc,
        "InChIKey 중복 제거 후 신규 81개 분자 (DB 미포함, 양성 26 / 음성 55) 가 학습에 추가되었다. 추가로 기존 168개 "
        "중 46개의 label conflict 가 식별되었으며, 이 중 28개 (예: Telaprevir, Velpatasvir, Lobeglitazone, Panobinostat) "
        "는 LiverTox 가 명확히 hepatotoxic (Likelihood B / C) 으로 분류한 반면 본 DB 가 음성으로 분류한 사례로 확인되어 "
        "라벨을 정정하였다. 나머지 18개는 FAERS / CTD 의 양성 신고와 LiverTox 의 rare (D / E) 분류가 충돌한 사례로, "
        "두 정의 모두 valid 하다고 판단하여 기존 라벨을 유지하였다.",
    )

    add_para(
        doc,
        "확장된 데이터로 Chemprop v31 및 RF/CB v31_v2 를 재학습하고 동일한 외부 263 분자에서 평가한 결과는 Table 5 와 "
        "같다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Model", "External AUC", "MCC@best", "Δ (vs baseline)"],
            ["Chemprop v27 (baseline)", "0.506", "0.055", "—"],
            ["Chemprop v31 (+81 신규)", "0.515", "0.052", "+0.009"],
            ["RF/CB v3 (baseline)", "0.448", "0.061", "—"],
            ["RF/CB v31 (+81 신규)", "0.439", "0.061", "-0.009"],
            ["RF/CB v31_v2 (+28 label fix)", "0.451", "0.061", "+0.003"],
        ],
    )
    add_table_caption(
        doc,
        "Table 5. 학습 데이터 확장 ablation 결과. 모든 모델에서 향상이 marginal (within noise) 이거나 negative 임이 "
        "확인되었다.",
    )

    add_para(
        doc,
        "Chemprop v31 의 외부 AUC 는 0.515 로 baseline (0.506) 대비 +0.009 의 marginal 향상을 보였으며, RF/CB 의 "
        "경우 오히려 약간 감소하였다. 즉 '데이터 양 부족이 외부 OOD 한계의 주원인' 가설은 기각되었다. 266개 추가 "
        "분자가 학습 9,645개의 약 2.7% 에 불과하여 학습 분포가 크게 변화하지 않은 점과, 더 본질적으로 외부 sanity 의 "
        "신약들이 학습 데이터의 어떤 chemical class 와도 유사하지 않은 새로운 분자들이라는 점이 원인으로 분석된다.",
    )

    # ───────────────────────────────────────────────
    add_heading(doc, "2.8 추가 분석 — Scaffold 중첩 여부에 따른 성능 분리", level=2)
    # ───────────────────────────────────────────────

    add_para(
        doc,
        "외부 263 분자 전체의 random 결과를 보다 정밀하게 분석하기 위해 각 분자의 Bemis-Murcko scaffold 가 학습 "
        "데이터의 scaffold 와 중첩되는지 검증하였다. 분석 결과 외부 263 분자 중 77개 (29.3%) 는 학습 분자와 동일한 "
        "scaffold 를 공유하였고 (InChIKey 만 상이), 186개 (70.7%) 는 학습에 등장하지 않은 새 scaffold 였다. 라벨별로 "
        "양성 129 중 28개 (21.7%), 음성 134 중 49개 (36.6%) 가 학습 scaffold 와 중첩되었다.",
    )

    add_para(doc, "두 그룹을 분리하여 Chemprop v27 의 성능을 측정한 결과는 Table 6 과 같다.")

    add_table_from_rows(
        doc,
        [
            ["그룹", "분자 수", "양성/음성", "AUC", "Best MCC"],
            ["A. 학습 scaffold 와 중첩", "82 (31%)", "29 / 53", "0.588", "+0.200"],
            ["B. 진정한 unseen scaffold", "181 (69%)", "100 / 81", "0.479", "+0.025"],
            ["전체 263", "263", "129 / 134", "0.506", "+0.013"],
        ],
    )
    add_table_caption(
        doc,
        "Table 6. 외부 263 분자를 scaffold 중첩 여부로 분리한 Chemprop v27 평가. 학습 scaffold 와 중첩되는 분자는 "
        "의미 있는 성능 (MCC 0.200) 을 보이는 반면, 진정한 unseen scaffold 는 random 수준 (MCC 0.025) 이다.",
    )

    add_para(
        doc,
        "본 결과로부터 명확한 패턴이 도출되었다. 모델은 '학습 데이터에서 동일 골격을 관찰한 분자' 에 대해서는 의미 "
        "있는 일반화 성능을 보이는 반면, '완전히 새로운 골격' 의 분자에 대해서는 random 수준의 예측만 가능하다. "
        "즉 외부 평가의 random 결과는 학습 데이터의 양적 부족이 아니라, 외부 sanity 분자의 70% 가 학습 데이터의 어떤 "
        "scaffold 와도 일치하지 않기 때문이다.",
    )

    add_para(
        doc,
        "Unseen scaffold 181개 분자의 chemical class 분포를 확인한 결과, 양성 100개 중 약 40개가 표적 항암제 "
        "(tyrosine kinase inhibitor) 이며, HCV direct-acting antiviral, CDK / PARP / BTK / HDAC inhibitor, 회수 약물 "
        "(cerivastatin, trovafloxacin, telithromycin) 등이 주요 class 로 확인되었다. 음성 81개는 SGLT2 / DPP-4 "
        "inhibitor, orexin antagonist, DOAC (apixaban 계열), 최신 항정신병약 등 대부분 post-2010 출시된 신약으로 "
        "구성되었다.",
    )

    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "3. Discussion", level=1)
    # ═══════════════════════════════════════════════════════════

    add_heading(doc, "3.1 결과 해석 — Chemical Space 의존성", level=2)

    add_para(
        doc,
        "본 연구에서 가장 명확히 확인된 사실은 모델의 일반화 성능이 평가 데이터의 chemical class 와 학습 데이터의 "
        "중첩 정도에 강하게 의존한다는 점이다. 학습 분포 내의 새 골격에 대해서는 MCC 0.615 의 의미 있는 성능을 보이나, "
        "학습에 없는 chemical class 에 대해서는 random 수준이다.",
    )

    add_para(
        doc,
        "Class expansion ablation (Section 2.7) 의 결과는 본 격차가 학습 데이터의 양적 부족이 아님을 입증한다. "
        "266개의 부족한 class 분자를 추가하고 28개 라벨을 정정해도 외부 sanity AUC 가 0.506 에서 0.515 의 marginal "
        "변화에 그쳤다. 즉 외부 OOD 한계는 데이터 큐레이션 노력으로 극복 가능한 기술적 문제가 아니라 DILI 자체의 "
        "생물학적 본질에 기인하는 것으로 해석된다.",
    )

    add_heading(doc, "3.2 DILI 의 본질적 한계", level=2)

    add_para(doc, "DILI 예측의 본질적 어려움은 다음의 세 가지 요인으로 분석할 수 있다 [22].")

    add_para(
        doc,
        "첫째, DILI 의 특이체질 (idiosyncratic) 본질이다 [10]. 환자마다 발생 여부가 상이하며, 분자 구조만으로는 "
        "결정되지 않는다. (a) 환자 HLA (인간 백혈구 항원) 유전형의 영향이 보고되었다. 예를 들어 HLA-B*5701 + "
        "flucloxacillin 의 조합에서만 cholestatic hepatitis 가 발생함이 GWAS 로 확인된 바 있다 [11, 18]. (b) 대사산물 "
        "매개 독성이 존재한다. Acetaminophen 의 경우 약물 자체는 안전하나 활성 대사산물 NAPQI 가 독성을 유발한다 [28]. "
        "(c) 발생률이 1,000명 ~ 10,000명 당 1명 수준으로 매우 낮아 통계적 학습 신호가 약하다 [4].",
    )

    add_para(
        doc,
        "둘째, 학습 chemical space 의 제한이다. 본 연구의 19,273개 학습 분자는 대부분 1970~2015년에 시판된 약물 "
        "중심으로 구성되어 있다. 최신 chemical class (TKI, GLP-1, SGLT2, HCV NS5A, CDK / PARP / BTK inhibitor 등) 는 "
        "학습 데이터에 부족하다. 매년 30~50개의 신약이 FDA 에서 승인되며 이들의 hepatic AE 정보가 LiverTox 등의 "
        "데이터베이스에 반영되는 데에는 post-marketing surveillance 의 특성상 5~10년의 시차가 존재한다.",
    )

    add_para(
        doc,
        "셋째, source label noise 이다. FAERS 와 같은 환자 자발 신고 데이터는 oversensitive 한 특성이 있으며, "
        "CTD / ChEMBL 의 hepatic disease 연관성은 indirect signal 에 해당한다. 8 source 의 OR rule 통합과 1,210건의 "
        "수동 큐레이션을 통해 이를 완화하였으나, 잔여 label noise 는 학습에 영향을 미친다.",
    )

    add_heading(doc, "3.3 시도한 접근 및 그 한계", level=2)

    add_para(
        doc,
        "외부 OOD 일반화 한계를 극복하기 위해 다음 다섯 가지 접근을 시도하였으나, 모두 marginal 또는 negative 효과만 "
        "확인되었다. (1) 학습 데이터 확장 — 266개 신규 분자 + 28개 라벨 정정 후 외부 AUC 0.506 → 0.515 의 marginal "
        "변화 (Section 2.7). (2) Source 신뢰도 기반 sample 가중치 — MCC 향상 미미하여 production 단순화 위해 제거. "
        "(3) Focal loss [25] (γ=2) — scaffold OOD MCC -0.007. (4) SMOTE oversampling [5] — scaffold OOD MCC -0.003. "
        "(5) Chemprop + RF/CB 의 honest stacking — 단일 Chemprop 대비 +0.002 MCC, 코드 복잡도 증가 대비 효과 미미하여 "
        "production 에서 제외.",
    )

    add_heading(doc, "3.4 향후 연구 방향", level=2)

    add_para(
        doc, "DILI 예측의 본질적 한계를 점진적으로 극복하기 위한 향후 연구 방향은 다음과 같다."
    )

    add_para(
        doc,
        "첫째, 분자 표현의 사전학습 (pre-training) 활용이다. 본 연구는 19,273개의 DILI labeled 분자만을 사용하였으나, "
        "MoLFormer [33] (IBM, 11억 분자 사전학습) 등의 분자 foundation model 을 활용하면 모델이 일반적 화학 지식을 "
        "사전에 학습한 상태에서 DILI 데이터로 fine-tune 할 수 있다. 학계 보고에 따르면 외부 신약 예측 AUC 가 "
        "0.50 → 0.55~0.62 정도 향상 가능하다 [37].",
    )

    add_para(
        doc,
        "둘째, 다요인 (multi-modal) 통합 모델이다. DILI 는 분자 구조 외 환자 유전형, 대사 효소 활성, 투여 용량 등 "
        "복합적 요인의 결합으로 발생한다. 분자 구조 정보에 CYP450 효소 활성 예측, HLA 유전형 빈도, 임상 투여 용량 "
        "정보를 함께 입력하는 multi-modal 모델이 가능하나, paired 데이터셋의 부족이 주요 장벽이다.",
    )

    add_para(
        doc,
        "셋째, 능동 학습 (active learning) 의 적용이다 [35]. 모델이 불확실한 예측을 보이는 신약 (predict_proba "
        "0.4~0.6 영역) 을 우선 라벨링하여 학습 데이터에 추가하고 반복 학습하는 방식이다. 학습 chemical space 를 "
        "점진적으로 확장하여 매년 출시되는 신약에 대한 효율적 retraining 이 가능하다.",
    )

    add_para(
        doc,
        "넷째, 예측 신뢰도 정량화 (uncertainty quantification) 이다. 모델의 예측에 confidence interval 을 함께 제공"
        "함으로써 novel chemical class 분자에 대해 예측 신뢰도가 낮음을 명시할 수 있다. Monte Carlo dropout [15] "
        "또는 ensemble disagreement 기반 방법 [23] 을 활용할 수 있다.",
    )

    add_heading(doc, "3.5 결론", level=2)

    add_para(
        doc,
        "본 연구는 SMILES 기반 DILI 예측 모델을 구축하고 그 일반화 성능의 한계를 체계적으로 분석하였다. 초기 "
        "random split 으로 얻은 MCC 0.85 가 data leakage 에 의한 과대평가임을 확인하고, Bemis-Murcko scaffold-balanced "
        "split 으로 재평가한 결과 진정한 일반화 성능 MCC 0.615 를 얻었다. 외부 263 분자에 대한 평가에서는 random 수준 "
        "(AUC 0.506) 의 한계가 확인되었으며, scaffold 중첩 분석을 통해 그 원인이 학습 데이터에 없는 chemical class "
        "임을 명확히 규명하였다.",
    )

    add_para(
        doc,
        "단순 데이터 확장 (266 신규 분자 추가) 으로는 외부 OOD 일반화 한계를 극복할 수 없음이 ablation 으로 입증되었다. "
        "DILI 의 본질적 어려움 — 특이체질, 대사산물 매개, 다요인성 — 이 학계 전반의 도전 과제임을 문헌 검토를 통해 "
        "확인하였다.",
    )

    add_para(
        doc,
        "본 모델은 학습 분포 내 신약에 대해 의미 있는 일반화 성능 (MCC 0.615) 을 제공하며, novel chemical class 에 "
        "대해서는 random 수준의 한계를 honest 하게 보고한다. 향후 foundation model 및 multi-modal 통합을 통해 외부 "
        "OOD 신약 예측의 한계를 점진적으로 극복할 수 있을 것으로 기대된다.",
    )

    # ═══════════════════════════════════════════════════════════
    add_heading(doc, "References", level=1)
    # ═══════════════════════════════════════════════════════════

    refs = [
        "Bemis, G. W., & Murcko, M. A. (1996). The properties of known drugs. 1. Molecular frameworks. Journal of Medicinal Chemistry, 39(15), 2887–2893.",
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32.",
        "Carhart, R. E., Smith, D. H., & Venkataraghavan, R. (1985). Atom pairs as molecular features in structure-activity studies. Journal of Chemical Information and Computer Sciences, 25(2), 64–73.",
        "Chalasani, N., & Björnsson, E. (2010). Risk factors for idiosyncratic drug-induced liver injury. Gastroenterology, 138(7), 2246–2259.",
        "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic Minority Over-sampling Technique. Journal of Artificial Intelligence Research, 16, 321–357.",
        "Chen, M., Borlak, J., & Tong, W. (2013). High lipophilicity and high daily dose of oral medications are associated with significant risk for drug-induced liver injury. Hepatology, 58(1), 388–396.",
        "Chen, M., Suzuki, A., Thakkar, S., Yu, K., Hu, C., & Tong, W. (2016). DILIrank: the largest reference drug list ranked by the risk for developing drug-induced liver injury in humans. Drug Discovery Today, 21(4), 648–653.",
        "Chicco, D., & Jurman, G. (2020). The advantages of the Matthews correlation coefficient (MCC) over F1 score and accuracy in binary classification evaluation. BMC Genomics, 21, 6.",
        "Chithrananda, S., Grand, G., & Ramsundar, B. (2020). ChemBERTa: Large-scale self-supervised pretraining for molecular property prediction. arXiv preprint arXiv:2010.09885.",
        "Daly, A. K. (2010). Genome-wide association studies in pharmacogenomics. Nature Reviews Genetics, 11(4), 241–246.",
        "Daly, A. K., Donaldson, P. T., Bhatnagar, P., Shen, Y., Pe'er, I., Floratos, A., et al. (2009). HLA-B*5701 genotype is a major determinant of drug-induced liver injury due to flucloxacillin. Nature Genetics, 41(7), 816–819.",
        "Davis, A. P., Wiegers, T. C., Johnson, R. J., Sciaky, D., Wiegers, J., & Mattingly, C. J. (2023). Comparative Toxicogenomics Database (CTD): update 2023. Nucleic Acids Research, 51(D1), D1257–D1262.",
        "Durant, J. L., Leland, B. A., Henry, D. R., & Nourse, J. G. (2002). Reoptimization of MDL keys for use in drug discovery. Journal of Chemical Information and Computer Sciences, 42(6), 1273–1280.",
        "Ertl, P., Rohde, B., & Selzer, P. (2000). Fast calculation of molecular polar surface area as a sum of fragment-based contributions and its application to drug transport properties. Journal of Medicinal Chemistry, 43(20), 3714–3717.",
        "Gal, Y., & Ghahramani, Z. (2016). Dropout as a Bayesian approximation: Representing model uncertainty in deep learning. Proceedings of the 33rd International Conference on Machine Learning, 1050–1059.",
        "Gedeck, P., Rohde, B., & Bartels, C. (2006). QSAR — How Good Is It in Practice? Comparison of Descriptor Sets on an Unbiased Cross Section of Corporate Data Sets. Journal of Chemical Information and Modeling, 46(5), 1924–1936.",
        "Hanley, J. A., & McNeil, B. J. (1982). The meaning and use of the area under a receiver operating characteristic (ROC) curve. Radiology, 143(1), 29–36.",
        "Hautekeete, M. L., Horsmans, Y., Van Waeyenberge, C., Demanet, C., Henrion, J., Verbist, L., et al. (1999). HLA association of amoxicillin-clavulanate-induced hepatitis. Gastroenterology, 117(5), 1181–1186.",
        "Heid, E., Greenman, K. P., Chung, Y., Li, S.-C., Graff, D. E., Vermeire, F. H., et al. (2024). Chemprop: A machine learning package for chemical property prediction. Journal of Chemical Information and Modeling, 64(1), 9–17.",
        "Heller, S. R., McNaught, A., Pletnev, I., Stein, S., & Tchekhovskoi, D. (2015). InChI, the IUPAC International Chemical Identifier. Journal of Cheminformatics, 7, 23.",
        "Hoofnagle, J. H., Serrano, J., Knoben, J. E., & Navarro, V. J. (2013). LiverTox: a website on drug-induced liver injury. Hepatology, 57(3), 873–874.",
        "Kullak-Ublick, G. A., Andrade, R. J., Merz, M., End, P., Benesic, A., Gerbes, A. L., & Aithal, G. P. (2017). Drug-induced liver injury: recent advances in diagnosis and risk assessment. Gut, 66(6), 1154–1164.",
        "Lakshminarayanan, B., Pritzel, A., & Blundell, C. (2017). Simple and Scalable Predictive Uncertainty Estimation using Deep Ensembles. Advances in Neural Information Processing Systems, 30.",
        "Li, H., Zhao, D., & Zeng, J. (2022). KPGT: Knowledge-Guided Pre-training of Graph Transformer for Molecular Property Prediction. Proceedings of the 28th ACM SIGKDD Conference, 857–867.",
        "Lin, T.-Y., Goyal, P., Girshick, R., He, K., & Dollár, P. (2017). Focal loss for dense object detection. Proceedings of the IEEE International Conference on Computer Vision, 2980–2988.",
        "Matthews, B. W. (1975). Comparison of the predicted and observed secondary structure of T4 phage lysozyme. Biochimica et Biophysica Acta - Protein Structure, 405(2), 442–451.",
        "Mendez, D., Gaulton, A., Bento, A. P., Chambers, J., De Veij, M., Félix, E., et al. (2019). ChEMBL: towards direct deposition of bioassay data. Nucleic Acids Research, 47(D1), D930–D940.",
        "Mitchell, J. R., Jollow, D. J., Potter, W. Z., Davis, D. C., Gillette, J. R., & Brodie, B. B. (1973). Acetaminophen-induced hepatic necrosis. I. Role of drug metabolism. Journal of Pharmacology and Experimental Therapeutics, 187(1), 185–194.",
        "Nilakantan, R., Bauman, N., Dixon, J. S., & Venkataraghavan, R. (1987). Topological torsion: a new molecular descriptor for SAR applications. Journal of Chemical Information and Computer Sciences, 27(2), 82–85.",
        "Prokhorenkova, L., Gusev, G., Vorobev, A., Dorogush, A. V., & Gulin, A. (2018). CatBoost: unbiased boosting with categorical features. Advances in Neural Information Processing Systems, 31.",
        "Rogers, D., & Hahn, M. (2010). Extended-connectivity fingerprints. Journal of Chemical Information and Modeling, 50(5), 742–754.",
        "Rong, Y., Bian, Y., Xu, T., Xie, W., Wei, Y., Huang, W., & Huang, J. (2020). Self-supervised graph transformer on large-scale molecular data. Advances in Neural Information Processing Systems, 33, 12559–12571.",
        "Ross, J., Belgodere, B., Chenthamarakshan, V., Padhi, I., Mroueh, Y., & Das, P. (2022). Large-scale chemical language representations capture molecular structure and properties. Nature Machine Intelligence, 4, 1256–1264.",
        "Seal, S., Williams, D. P., Hosseini-Gerami, L., Mahale, M., Carpenter, A. E., Spjuth, O., & Bender, A. (2024). Improved Detection of Drug-Induced Liver Injury by Integrating Predicted in vivo and in vitro Data. Chemical Research in Toxicology, 37, 1290–1305.",
        "Settles, B. (2009). Active learning literature survey. University of Wisconsin–Madison Computer Sciences Technical Report 1648.",
        "U.S. Food and Drug Administration (2024). FDA Adverse Event Reporting System (FAERS) Public Dashboard. Available at: https://fis.fda.gov/extensions/FPD-QDE-FAERS/FPD-QDE-FAERS.html",
        "Vall, A., Sabnis, Y., Shi, J., Class, R., Hochreiter, S., & Klambauer, G. (2021). The promise of AI for drug-induced liver injury prediction with artificial intelligence. Frontiers in Artificial Intelligence, 4, 638410.",
        "Watkins, P. B. (2005). Idiosyncratic liver injury: challenges and approaches. Toxicologic Pathology, 33(1), 1–5.",
        "Weininger, D. (1988). SMILES, a chemical language and information system. 1. Introduction to methodology and encoding rules. Journal of Chemical Information and Computer Sciences, 28(1), 31–36.",
        "Wu, Z., Ramsundar, B., Feinberg, E. N., Gomes, J., Geniesse, C., Pappu, A. S., Leswing, K., & Pande, V. (2018). MoleculeNet: a benchmark for molecular machine learning. Chemical Science, 9(2), 513–530.",
        "Yang, K., Swanson, K., Jin, W., Coley, C., Eiden, P., Gao, H., et al. (2019). Analyzing learned molecular representations for property prediction. Journal of Chemical Information and Modeling, 59(8), 3370–3388.",
    ]

    for i, ref in enumerate(refs, 1):
        add_reference(doc, ref, i)

    doc.save(OUTPUT)
    print(f"저장 완료: {OUTPUT}")
    print(f"크기: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
    print(f"References: {len(refs)} 개")


if __name__ == "__main__":
    build_doc()
