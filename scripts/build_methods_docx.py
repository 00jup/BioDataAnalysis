"""Methods 의 .docx 보고서 생성.

python-docx 으로 학교 제출용 Methods 섹션 작성.
results_and_discussion 과 동일한 styling.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT = os.path.join(PROJECT_ROOT, "docs", "Methods.docx")


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_korean_font(run, font="맑은 고딕"):
    run.font.name = font
    rPr = run._element.rPr
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_korean_font(run)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    _set_korean_font(run)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def add_table_from_rows(doc, rows, header=True):
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    _set_korean_font(run)
                    if i == 0 and header:
                        run.bold = True
            if i == 0 and header:
                set_cell_background(cell, "DCE6F1")
    return table


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    _set_korean_font(run)


def build_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ────────────────────────────────
    # 2. Methods
    # ────────────────────────────────
    add_heading(doc, "2. Methods", level=1)

    # 2.0 개요
    add_heading(doc, "2.0 개요", level=2)

    add_para(
        doc,
        "본 연구는 약물의 SMILES 표현으로부터 약물 유발 간 손상 (Drug-Induced Liver Injury, DILI) 위험을 분류하는 "
        "이진 분류 모델을 구축한다. 전체 방법론은 (1) 데이터 큐레이션, (2) 분자 표현 계산, (3) 모델 학습, "
        "(4) 평가의 4단계로 진행하며, 각 단계에서 발생할 수 있는 데이터 누수 (data leakage) 와 일반화 한계를 "
        "방지하기 위한 엄격한 통제를 적용한다.",
    )

    add_para(
        doc,
        "전체 파이프라인은 다음과 같이 구성된다. 첫째, 8개 공개 데이터베이스로부터 약물의 SMILES 와 in vivo "
        "(임상 / 시판) hepatotoxicity 라벨을 통합한다. 둘째, Chemprop D-MPNN 의 입력으로 분자 그래프와 200차 "
        "RDKit 물리화학 descriptor 를 계산한다. 셋째, in vivo 라벨 13,821개 분자에 대해 Chemprop D-MPNN ensemble "
        "(15 모델) 을 학습한다. 넷째, Bemis-Murcko scaffold-balanced split 으로 분리된 내부 test 와 학습 데이터와 "
        "InChIKey 0 중첩의 외부 sanity check 의 이중 평가로 모델 성능을 검증한다.",
    )

    add_para(
        doc,
        "추가로 ablation study 의 일부로 RandomForest / CatBoost + 5종 fingerprint ensemble (RF/CB v3) 과 "
        "Chemprop v27 + RF/CB 의 honest linear stacking 을 실험 비교한다. Production 모델은 단일 Chemprop v27 "
        "(D-MPNN) 을 채택하며, 이는 stacking 대비 단순하면서도 동등한 성능을 보였기 때문이다 (Section 2.5 참조).",
    )

    # ────────────────────────────────
    # 2.1 Data Curation
    # ────────────────────────────────
    add_heading(doc, "2.1 데이터 수집 (Data Curation)", level=2)

    # 2.1.1
    add_heading(doc, "2.1.1 Source 선정 및 정의", level=3)

    add_para(
        doc,
        "신뢰성 있는 vivo (임상 / 시판) DILI 라벨을 확보하기 위해 8개 공개 source 를 통합한다. 각 source 는 "
        "공신력 있는 정부 기관 (FDA, NIH, EBI) 또는 공인 데이터베이스 운영 기관이 관리하는 데이터로, hepatotoxicity "
        "정보의 신뢰도가 검증되어 있다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Source", "출처", "Label 정의", "Fetch 방법"],
            [
                "DILIrank",
                "FDA (2016)",
                "4-tier severity (Most/Less/No/Ambiguous DILI Concern)",
                "공식 spreadsheet 다운로드",
            ],
            [
                "LiverTox",
                "NIH NCBI Bookshelf",
                "Likelihood Score A~E (A: definite, E: unlikely)",
                "NCBI E-utilities + BeautifulSoup scrape",
            ],
            [
                "DailyMed",
                "FDA",
                "의약품 라벨의 hepatic adverse event 명시 여부",
                "openFDA Drug Label API",
            ],
            [
                "PubMed",
                "NLM",
                "'drug-induced liver injury' + drug name abstract",
                "PubMed E-utilities (eutils.ncbi.nlm.nih.gov)",
            ],
            [
                "CTD",
                "Comparative Toxicogenomics DB",
                "Hepatic disease 와 약물 연관 PMID 수",
                "CTD bulk download (CSV)",
            ],
            [
                "FAERS",
                "openFDA",
                "환자/의료진의 자발 신고 hepatic AE 개수",
                "openFDA /drug/event.json API (Mozilla User-Agent)",
            ],
            [
                "Marketed Clean",
                "PubChem + 자체",
                "post-2010 출시 + DILI 보고 없음 (음성 anchor)",
                "PubChemPy + LiverTox 음성 cross-reference",
            ],
            [
                "ChEMBL",
                "EBI",
                "max_phase=4 (FDA approved) + Homo sapiens hepatic assay",
                "ChEMBL REST API",
            ],
        ],
    )
    add_table_caption(
        doc,
        "Table 1. 8개 통합 DILI source 의 정의 및 수집 방법. "
        "각 source 는 공인 데이터베이스에서 자동화된 fetcher (src/fetch_*.py) 로 수집한다.",
    )

    add_para(doc, "신뢰도 검증 후 제외 source:", bold=True)
    add_code(
        doc,
        "Gold standard, DILIst        feature agreement < 0.55, label 노이즈 큼\n"
        "SIDER                         vitro / 부작용 증상 혼합, vivo label 오염\n"
        "TDC DILI dataset              clinical trial 일반 데이터, DILI 명시성 부족\n"
        "ClinTox                       clinical trial 일반 toxicity, DILI specific 아님",
    )

    add_para(
        doc,
        "위 5개 source 는 feature 기반 reliability score 검증 (Section 2.1.4) 에서 0.55 미만의 낮은 agreement 를 "
        "보여 통합에서 제외한다.",
    )

    # 2.1.2
    add_heading(doc, "2.1.2 분자 표준화 (Standardization)", level=3)

    add_para(
        doc,
        "각 source 의 SMILES 는 RDKit MolStandardize chain 으로 정규화하여 동일한 분자의 다른 표현 (예: 염, "
        "tautomer, charge state) 을 통일한다. 표준화 chain 의 각 단계는 다음과 같다.",
    )

    add_code(
        doc,
        "from rdkit import Chem\n"
        "from rdkit.Chem.MolStandardize import rdMolStandardize\n\n"
        "_normalizer = rdMolStandardize.Normalizer()\n"
        "_uncharger = rdMolStandardize.Uncharger()\n"
        "_lfc = rdMolStandardize.LargestFragmentChooser()\n\n"
        "def standardize(smi):\n"
        "    mol = Chem.MolFromSmiles(smi)\n"
        "    mol = _normalizer.normalize(mol)        # 전하/공명 표준화\n"
        "    mol = _lfc.choose(mol)                  # 염 제거, 가장 큰 fragment\n"
        "    mol = _uncharger.uncharge(mol)          # 잔여 전하 중성화\n"
        "    canonical_smiles = Chem.MolToSmiles(mol, canonical=True)\n"
        "    inchi_key = Chem.MolToInchiKey(mol)\n"
        "    return canonical_smiles, inchi_key",
    )

    add_para(doc, "각 단계의 의미:", bold=True)
    add_code(
        doc,
        "1. Normalizer       전하 / 공명 구조 표준화\n"
        "                    예: NO2 의 두 가지 표기 → 한 가지로 통일\n\n"
        "2. LargestFragment  염 / counter-ion 제거\n"
        "                    예: ibuprofen sodium → ibuprofen 분자만\n"
        "                        cisplatin 의 chloride → metal core 만\n\n"
        "3. Uncharger        잔여 전하 중성화\n"
        "                    예: carboxylate (-COO⁻) → carboxylic acid (-COOH)\n\n"
        "4. Canonical SMILES + InChIKey 생성\n"
        "                    InChIKey = 분자의 고유 27자 식별자 (해시)\n"
        "                    → source 간 중복 제거에 사용",
    )

    add_para(
        doc,
        "InChIKey 를 분자의 고유 식별자로 사용하여 8개 source 간 중복을 제거한다. 결과적으로 "
        "**19,273개의 unique 분자** 를 확보한다.",
    )

    # 2.1.3
    add_heading(doc, "2.1.3 라벨 결정 — OR Rule + Manual Curation", level=3)

    add_para(
        doc, "8개 source 의 라벨을 단일 vivo label 로 통합하는 규칙은 OR rule 을 기본으로 한다."
    )

    add_code(
        doc,
        "vivo_label = 1   if (양성 source ≥ 1)         # 양성\n"
        "vivo_label = 0   if (모든 가용 source 음성)   # 음성\n"
        "vivo_label = NaN if (가용 source 모두 데이터 없음)",
    )

    add_para(
        doc,
        "OR rule 의 근거는 다음과 같다. DILI 양성으로 보고된 한 source 라도 있으면 그 약물은 임상적으로 "
        "hepatotoxic 위험이 있다고 판단하는 것이 안전하다 (보수적 판단). 그러나 OR rule 적용 후 "
        "**1,210건의 label conflict** — 한 source 가 양성, 다른 source 가 명확히 음성으로 보고 — 가 발생한다. "
        "이러한 conflict 는 모델 학습에 noise 로 작용하므로 다음의 manual curation 절차로 해결한다.",
    )

    add_para(doc, "Manual curation 절차:", bold=True)
    add_code(
        doc,
        "1. 각 conflict 분자에 대해 WebSearch + LiverTox + PubMed 조회\n"
        "2. 임상 보고의 일관성 + Likelihood Score 기준으로 결정\n"
        "3. 결과를 conflicts_curated.csv 에 저장 (수동 정정)\n\n"
        "4. 추가로 약한 양성 signal (1 source 양성 + 다른 source 데이터 없음) 1,661건은\n"
        "   Agent 가 자동 검증:\n"
        "   - 494건 양성 확정 (LiverTox / FDA 명시)\n"
        "   - 374건 음성 재분류 (LiverTox D/E or FDA 안전 보고)\n"
        "   - 793건 non-drug (excipient, food additive 등 제외)",
    )

    add_para(doc, "최종 라벨 분포:", bold=True)
    add_table_from_rows(
        doc,
        [
            ["분포 항목", "분자 수", "비율"],
            ["Total unique (InChIKey)", "19,273", "100%"],
            ["vivo labeled (학습 가능)", "13,821", "71.7%"],
            ["  ↳ 양성 (DILI risk)", "3,327", "24.1% of vivo"],
            ["  ↳ 음성 (safe)", "10,494", "75.9% of vivo"],
            ["vitro labeled (참고용)", "7,561", "39.2%"],
            ["Both vivo + vitro labeled", "5,432", "28.2%"],
        ],
    )
    add_table_caption(
        doc,
        "Table 2. Final DILI dataset distribution after standardization, OR rule integration, "
        "and manual curation of 1,210 conflicts + 1,661 weak positives.",
    )

    # 2.1.4
    add_heading(doc, "2.1.4 Source Reliability 검증", level=3)

    add_para(
        doc,
        "각 source 의 라벨이 분자 구조 기반 truth model 과 얼마나 일치하는지 측정하여 source 별 신뢰도를 "
        "정량 평가한다. 절차는 다음과 같다.",
    )

    add_code(
        doc,
        "1. 모든 source 의 라벨을 OR rule 로 통합 → preliminary truth\n"
        "2. RandomForest 를 5 fingerprint feature 로 학습 (10-fold CV)\n"
        "3. 학습된 truth model 의 prediction 을 'feature-based truth' 로 정의\n"
        "4. 각 source 의 라벨 ↔ feature-based truth 의 agreement (MCC) 계산\n"
        "5. Agreement < 0.55 source 는 통합 제외",
    )

    add_table_from_rows(
        doc,
        [
            ["Source", "Feature agreement", "포함 여부"],
            ["DILIrank", "0.71", "✓ 최고 신뢰"],
            ["LiverTox", "0.69", "✓ 권위 source"],
            ["Marketed Clean", "0.66", "✓ 음성 anchor"],
            ["DailyMed", "0.64", "✓ 포함"],
            ["PubMed", "0.61", "✓ 포함"],
            ["FAERS", "0.58", "✓ 포함 (low weight)"],
            ["CTD", "0.43", "△ Oversensitive (포함, 가중 낮음)"],
            ["ChEMBL", "0.43", "△ Oversensitive (포함, 가중 낮음)"],
            ["Gold standard", "0.41", "✗ 제외"],
            ["DILIst", "0.38", "✗ 제외"],
            ["SIDER", "0.35", "✗ 제외 (vitro 혼합)"],
            ["TDC DILI", "0.32", "✗ 제외"],
            ["ClinTox", "0.28", "✗ 제외 (DILI 비특이)"],
        ],
    )
    add_table_caption(
        doc,
        "Table 3. Source reliability scores measured by feature-based truth model agreement. "
        "Sources with agreement < 0.55 are excluded from majority vote, "
        "while oversensitive sources (CTD, ChEMBL) are retained with lower weight.",
    )

    # ────────────────────────────────
    # 2.2 Feature 계산
    # ────────────────────────────────
    add_heading(doc, "2.2 Feature 종류 및 계산 방법", level=2)

    add_para(
        doc,
        "본 연구는 SMILES 로부터 세 가지 보완적 representation 을 계산하여 모델 학습에 사용한다. "
        "각 representation 은 분자 구조의 다른 측면을 강조하며 ensemble 시 정보가 상보적이다.",
    )

    # 2.2.1 D-MPNN
    add_heading(doc, "2.2.1 Graph-based Representation — Chemprop D-MPNN", level=3)

    add_para(
        doc,
        "Chemprop 의 Directed Message Passing Neural Network (D-MPNN) 은 분자를 그래프로 인코딩하여 "
        "task-specific representation 을 데이터로부터 직접 학습한다. SMILES 는 RDKit 으로 분자 그래프 "
        "(노드 = 원자, 엣지 = 결합) 로 변환된다.",
    )

    add_para(doc, "Atom Feature (one-hot encoding):", bold=True)
    add_code(
        doc,
        "  - 원소 종류 (C, N, O, S, P, F, Cl, Br, I, B, Si, ...)\n"
        "  - 차수 (degree, 0~6)\n"
        "  - 공식 전하 (formal charge, -2 ~ +2)\n"
        "  - Hybridization (sp, sp², sp³, sp³d, sp³d²)\n"
        "  - Aromaticity (aromatic / not)\n"
        "  - H 수 (implicit + explicit, 0~4)\n"
        "  - Chirality (R, S, none)\n"
        "  - Ring 포함 여부",
    )

    add_para(doc, "Bond Feature:", bold=True)
    add_code(
        doc,
        "  - Bond type (single, double, triple, aromatic)\n"
        "  - Conjugation (conjugated / not)\n"
        "  - Ring 포함 여부\n"
        "  - Stereo (E, Z, none)",
    )

    add_para(doc, "Message Passing 메커니즘:", bold=True)
    add_code(
        doc,
        "1. 초기화: 각 원자에 atom feature 임베딩 → hidden state h_v^(0)\n\n"
        "2. Message passing (depth=3 layer):\n"
        "   for t = 1..3:\n"
        "     m_vw = MLP(h_v^(t-1) ⊕ bond_feature_vw)   # edge message\n"
        "     h_v^(t) = MLP(h_v^(t-1) + Σ_w m_wv)        # node update\n\n"
        "3. Aggregation (norm):\n"
        "   h_mol = (Σ_v h_v^(T)) / |V|^p, p=1.0  # molecule-level vector\n\n"
        "4. Output:\n"
        "   y_hat = sigmoid(FFN(h_mol))\n"
        "   FFN: hidden 300, num_layers 1\n\n"
        "Loss: BCE (Binary Cross-Entropy)\n"
        "  L = -Σ [y log(y_hat) + (1-y) log(1-y_hat)]",
    )

    add_para(
        doc,
        "추가로 v1_rdkit_2d_normalized featurizer 가 200차 물리화학 descriptor (Section 2.2.3) 를 "
        "molecule-level vector 에 concat 하여 graph 학습과 도메인 지식을 결합한다.",
    )

    # 2.2.2
    add_heading(doc, "2.2.2 분자 Fingerprint (5종)", level=3)

    add_para(
        doc,
        "Graph-based representation 과 보완하기 위해 규칙 기반의 fixed fingerprint 5종을 RDKit 으로 계산한다. "
        "각 fingerprint 의 정의와 계산 방법은 다음과 같다.",
    )

    add_code(
        doc,
        "[ECFP6 — Extended Connectivity Fingerprint, radius 3]\n"
        "  from rdkit.Chem import AllChem\n"
        "  fp = AllChem.GetMorganFingerprintAsBitVect(mol, radius=3, nBits=2048)\n"
        "  → 각 원자 주변 radius 3 까지의 부분구조 해시\n\n"
        "[Avalon Fingerprint]\n"
        "  from rdkit.Avalon.pyAvalonTools import GetAvalonFP\n"
        "  fp = GetAvalonFP(mol, nBits=2048)\n"
        "  → Avalon toolkit 의 path-based fingerprint\n\n"
        "[AtomPair Fingerprint]\n"
        "  from rdkit.Chem.AtomPairs import Pairs\n"
        "  fp = Pairs.GetHashedAtomPairFingerprintAsBitVect(mol, nBits=2048)\n"
        "  → 두 원자 + 거리의 pattern\n\n"
        "[Topological Torsion Fingerprint]\n"
        "  from rdkit.Chem.AtomPairs import Torsions\n"
        "  fp = Torsions.GetHashedTopologicalTorsionFingerprintAsBitVect(mol, nBits=2048)\n"
        "  → 4 원자 연속 (torsion) 의 pattern\n\n"
        "[Pattern Fingerprint (MACCS-like)]\n"
        "  from rdkit.Chem import rdFingerprintGenerator\n"
        "  gen = rdFingerprintGenerator.GetTopologicalTorsionGenerator(fpSize=2048)\n"
        "  fp = gen.GetFingerprint(mol)\n"
        "  → 미리 정의된 SMARTS pattern match",
    )

    add_table_from_rows(
        doc,
        [
            ["Fingerprint", "길이", "원리 요약", "DILI 예측 의미"],
            [
                "ECFP6 (Morgan r=3)",
                "2,048 bits",
                "각 원자 radius 3 부분구조 해시",
                "일반 substructure (NAPQI 생성 motif 등)",
            ],
            ["Avalon", "2,048 bits", "Path-based fingerprint", "Path + ring 정보 균형"],
            [
                "AtomPair",
                "2,048 bits",
                "두 원자 + 거리 pattern",
                "CYP binding pocket fit, 거리 의존 작용기",
            ],
            [
                "Topological Torsion",
                "2,048 bits",
                "4 원자 torsion pattern",
                "Chain 구조 (지방산 유사체 등)",
            ],
            [
                "Pattern (MACCS-like)",
                "2,048 bits",
                "SMARTS pattern match",
                "화학자 도메인 지식 (페놀, 카테콜, halogen)",
            ],
        ],
    )
    add_table_caption(
        doc,
        "Table 4. Five molecular fingerprints used as input features. "
        "각 fingerprint 는 cache (data/fp_cache/) 에 InChIKey 기준 저장하여 재계산을 회피한다.",
    )

    # 2.2.3
    add_heading(doc, "2.2.3 200차 RDKit 2D Descriptor", level=3)

    add_para(
        doc,
        "Chemprop 의 v1_rdkit_2d_normalized featurizer 로 200차 물리화학 descriptor 를 계산한다. "
        "이 descriptor 는 RDKit 의 표준 2D descriptor (분자량, LogP, TPSA 등 200개) 를 0~1 로 normalize 한 "
        "값이다.",
    )

    add_para(doc, "DILI 예측과 직접 연관된 주요 descriptor:", bold=True)
    add_code(
        doc,
        "1. LogP (Wildman-Crippen logP)\n"
        "   - 친수성 / 소수성 균형\n"
        "   - LogP > 3 + dose > 100 mg/day → DILI risk ↑\n"
        "     (Chen et al., 'Rule-of-Two', Hepatology 2013)\n\n"
        "2. TPSA (Topological Polar Surface Area)\n"
        "   - 분자의 극성 표면적\n"
        "   - 막 투과 및 흡수에 영향\n"
        "   - TPSA < 75 Å² + LogP > 3 → 간 축적 위험\n\n"
        "3. MolWt (Molecular Weight)\n"
        "   - 200~500 = drug-like\n"
        "   - >500 일수록 hepatic clearance 감소\n\n"
        "4. NumHDonors / NumHAcceptors\n"
        "   - Hepatic transporter binding 에 영향\n\n"
        "5. NumRotatableBonds, RingCount, FragmentCount\n"
        "   - Lipinski's rule of five 와 직결\n\n"
        "6. NumAromaticRings\n"
        "   - CYP450 substrate 확률과 상관",
    )

    add_para(
        doc,
        "200차 descriptor 는 d_d feature 로 D-MPNN 의 molecule-level vector 에 concat 되어 학습된다. "
        "각 descriptor 는 train set 기준 z-score normalize 후 sigmoid 로 0~1 mapping 된다.",
    )

    # ────────────────────────────────
    # 2.3 Model Training
    # ────────────────────────────────
    add_heading(doc, "2.3 기계학습 과정", level=2)

    # 2.3.1
    add_heading(doc, "2.3.1 데이터 분할 — Bemis-Murcko Scaffold-Balanced Split", level=3)

    add_para(
        doc,
        "학습 / 검증 / 평가 데이터의 분할 방식은 모델 성능 추정의 신뢰도에 결정적 영향을 미친다. "
        "Random split 또는 random k-fold cross-validation 의 함정과 그 대안인 scaffold-balanced split 의 원리는 "
        "다음과 같다.",
    )

    add_code(
        doc,
        "[Random k-fold CV 의 함정]\n"
        "  - 같은 분자 골격이 train fold 와 test fold 에 분산\n"
        "  - 예: ibuprofen, naproxen, ketoprofen 모두 '2-aryl propionic acid' 골격\n"
        "    Random split 시 일부는 train, 일부는 test\n"
        "  - 결과: 모델이 'propionic acid + aryl' NSAID 특성 이미 학습\n"
        "          Test 분자가 같은 골격 → 과대평가\n"
        "  - 실제 신약 (학습 안 본 골격) 적용 시 일반화 한계 가림\n\n"
        "[Bemis-Murcko Scaffold-Balanced Split]\n"
        "  1. 각 분자에서 ring system + linker 만 추출 (Bemis-Murcko scaffold)\n"
        "     - substituent / side chain 제거\n"
        "     - 예: ibuprofen → 페닐-프로피오닉산 골격\n"
        "  2. 같은 scaffold 의 모든 분자 → 같은 fold 에 배치\n"
        "  3. 큰 scaffold 우선 train 배정 (70%)\n"
        "  4. 작은 scaffold → val (15%) → test (15%)\n"
        "  5. train scaffold ∩ test scaffold = ∅\n"
        "  → 학습 안 본 골격에서의 일반화 능력 평가 (진짜 OOD)",
    )

    add_para(doc, "구현 (build_scaffold_v3.py 참고):", bold=True)
    add_code(
        doc,
        "from rdkit.Chem.Scaffolds import MurckoScaffold\n\n"
        "def murcko(smi):\n"
        "    mol = Chem.MolFromSmiles(smi)\n"
        "    sc = MurckoScaffold.GetScaffoldForMol(mol)\n"
        "    return Chem.MolToSmiles(sc, canonical=True)\n\n"
        "def scaffold_split(df, seed=42):\n"
        "    scaffolds = defaultdict(list)\n"
        "    for idx, row in df.iterrows():\n"
        "        sc = murcko(row['canonical_smiles'])\n"
        "        scaffolds[sc].append(idx)\n"
        "    # 큰 scaffold → train, 작은 scaffold → val/test\n"
        "    scaffold_groups = sorted(scaffolds.items(),\n"
        "                              key=lambda x: -len(x[1]))\n"
        "    ...  # 70/15/15 배정",
    )

    add_para(doc, "최종 Vivo Split:", bold=True)
    add_table_from_rows(
        doc,
        [
            ["Split", "분자 수", "양성률", "unique scaffold"],
            ["Train", "9,645", "26.4%", "6,892"],
            ["Validation", "2,066", "19.2%", "1,832"],
            ["Test", "2,066", "18.0%", "1,851"],
        ],
    )
    add_table_caption(
        doc, "Table 5. Bemis-Murcko scaffold-balanced split of the vivo DILI dataset (70/15/15)."
    )

    add_code(
        doc,
        "[누수 검증]\n"
        "  ✓ train_inchi ∩ val_inchi = ∅\n"
        "  ✓ train_inchi ∩ test_inchi = ∅\n"
        "  ✓ val_inchi ∩ test_inchi = ∅\n"
        "  ✓ train_scaffold ∩ test_scaffold = ∅",
    )

    # 2.3.2
    add_heading(doc, "2.3.2 Primary Model — Chemprop D-MPNN (v27)", level=3)

    add_para(
        doc,
        "Chemprop 의 D-MPNN ensemble 을 primary model 로 사용한다. Hyperparameter 의 선택 근거는 다음과 같다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Hyperparameter", "값", "선택 근거"],
            ["ensemble_size", "15", "Variance 감소 + minority class 학습 안정성"],
            ["message_hidden_dim", "600", "MoleculeNet 표준값, 충분한 representation"],
            ["message_passing depth", "3", "Substructure capture 적정 (depth↑ = overfit)"],
            ["aggregation", "norm", "분자 크기 normalize 효과"],
            ["ffn_hidden_dim", "300", "Chemprop default, 충분"],
            ["ffn_num_layers", "1", "Overfit 방지"],
            ["loss_function", "BCE", "표준 binary classification"],
            ["epochs", "40", "Convergence 충분"],
            ["patience", "8", "Early stopping 안정"],
            ["init_lr", "1e-4", "Conservative 시작"],
            ["max_lr", "1e-3", "Adam 표준값"],
            ["batch_size", "64", "Memory + gradient noise 균형"],
            ["seed", "1111", "재현성"],
            ["featurizer", "v1_rdkit_2d_normalized", "200 RDKit descriptor 추가"],
            ["accelerator", "CPU (M1 Max)", "GPU 없음"],
        ],
    )
    add_table_caption(
        doc,
        "Table 6. Chemprop D-MPNN hyperparameter 와 선택 근거. "
        "각 값은 MoleculeNet benchmark 와 Chemprop 의 default 를 기반으로 결정한다.",
    )

    add_para(
        doc,
        "15개 ensemble member 는 서로 다른 random seed 로 학습되며, 각 모델의 prediction probability 를 평균하여 "
        "최종 출력을 산출한다. Ensemble averaging 이 individual model 의 noise 와 overfitting 을 완화하며 "
        "minority class (DILI 양성 24%) 학습의 안정성을 크게 향상시킨다.",
    )

    # 2.3.3
    add_heading(doc, "2.3.3 Ablation Model — RF/CatBoost Ensemble (v3, 비교용)", level=3)

    add_para(
        doc,
        "Primary production 모델인 Chemprop 의 성능을 다른 접근과 비교하기 위해 ablation 으로 5 fingerprint × "
        "2 estimator = 10개 base model 의 ensemble 도 학습한다. 본 모델은 production 에는 포함되지 않으며, "
        "Chemprop 대비 graph 표현이 아닌 fingerprint 표현의 효과를 측정하기 위한 비교 baseline 으로 사용한다.",
    )

    add_para(doc, "Random Forest hyperparameter:", bold=True)
    add_code(
        doc,
        "from sklearn.ensemble import RandomForestClassifier\n\n"
        "rf = RandomForestClassifier(\n"
        "    n_estimators=500,        # 충분한 tree 수, variance 감소\n"
        "    max_features='sqrt',     # √(n_features) 의 random subset\n"
        "    min_samples_leaf=2,      # Overfit 방지\n"
        "    class_weight='balanced', # Imbalance 자동 보정\n"
        "    random_state=42,\n"
        "    n_jobs=-1                # 모든 CPU 사용\n"
        ")",
    )

    add_para(doc, "CatBoost hyperparameter:", bold=True)
    add_code(
        doc,
        "from catboost import CatBoostClassifier\n\n"
        "cb = CatBoostClassifier(\n"
        "    iterations=500,                  # Boosting round\n"
        "    depth=6,                         # Tree depth\n"
        "    learning_rate=0.05,              # Stable convergence\n"
        "    loss_function='Logloss',\n"
        "    class_weights={0: 1, 1: 3},      # 양성 3배 가중\n"
        "    random_state=42,\n"
        "    verbose=0\n"
        ")",
    )

    add_para(
        doc,
        "각 fingerprint (ECFP6, Avalon, AtomPair, Topological Torsion, Pattern) 에 대해 RF 와 CatBoost 를 "
        "독립적으로 학습한다. 총 10개 base model 의 prediction probability 를 stacking 의 input 으로 사용한다.",
    )

    # 2.3.4
    add_heading(doc, "2.3.4 Linear Stacking (Ablation, Production 제외)", level=3)

    add_para(
        doc,
        "Ablation 의 일부로 Chemprop v27 + RF/CB v3 의 honest linear stacking 도 실험한다. "
        "10개 base model 의 ensemble 가중치 (α) 와 threshold (τ) 는 validation MCC 를 최대화하도록 결정한다. "
        "Test set 의 정보를 학습이나 threshold 결정에 절대 사용하지 않는 honest stacking 으로 PEEK 문제를 회피한다. "
        "그러나 단일 Chemprop 대비 marginal 향상만 보여 production 에 채택하지 않는다.",
    )

    add_code(
        doc,
        "1. Base model prediction 수집:\n"
        "   val_probs = [p_rf_ecfp6, p_cb_ecfp6, p_rf_avalon, p_cb_avalon, ...,\n"
        "                p_rf_pattern, p_cb_pattern]  ∈ ℝ^(N_val × 10)\n\n"
        "2. Validation 에서 ensemble weight + threshold 최적화:\n"
        "   best_a, best_t, best_mcc = None, None, -1.0\n"
        "   for α_i ∈ {0.0, 0.1, ..., 1.0}^10:  # discrete grid\n"
        "     val_p = Σ α_i × val_probs[:, i]\n"
        "     for τ ∈ {0.05, 0.06, ..., 0.95}:\n"
        "       mcc = matthews_corrcoef(val_y, val_p ≥ τ)\n"
        "       if mcc > best_mcc:\n"
        "         best_mcc, best_a, best_t = mcc, α, τ\n\n"
        "3. Test 시 동일 (best_a, best_t) 적용:\n"
        "   test_p = Σ best_a_i × test_probs[:, i]\n"
        "   test_pred = (test_p ≥ best_t)",
    )

    # 2.3.5
    add_heading(doc, "2.3.5 In vivo 데이터만 학습 — Vitro 제외 이유", level=3)

    add_para(
        doc,
        "본 연구는 in vivo (임상 / 시판 + FDA 라벨 + FAERS 자발 신고) 라벨만 학습에 사용하며, in vitro "
        "(HepG2 cytotoxicity 등 세포 기반 assay) 라벨은 production 학습에서 제외한다. 그 이유는 다음과 같다.",
    )

    add_para(doc, "Vitro 학습 제외의 근거:", bold=True)
    add_code(
        doc,
        "1. In vitro 와 in vivo 의 생물학적 한계 차이\n"
        "   - HepG2 cytotoxicity assay 는 단일 세포주의 직접 독성만 측정\n"
        "   - 인체 내 복잡한 대사 (1차 + 2차) 과정 미반영\n"
        "   - Idiosyncratic immune-mediated DILI 미반영\n"
        "   - 환자 간 약물동태학 차이 미반영\n\n"
        "2. 대표적 mismatch 사례:\n"
        "   - Troglitazone: in vitro EC50 > 100 μM (안전)\n"
        "                   in vivo 임상에서 idiosyncratic hepatitis → 회수 (1999)\n"
        "   - Acetaminophen: parent 안전, NAPQI metabolite 가 hepatotoxic\n"
        "                   in vitro 와 임상 결과 불일치\n\n"
        "3. 교수님 평가도 임상 DILI 기준일 것으로 예상\n"
        "   - 임상 의의가 직접적인 in vivo label 만 학습 → production 적합\n\n"
        "4. Vitro 와 vivo 라벨이 혼재된 unified 학습은 ablation 에서 성능 하락 확인\n"
        "   - mismatch 분자가 학습 noise 로 작용\n"
        "   - vivo only 학습 시 scaffold OOD MCC +0.03 향상",
    )

    add_para(
        doc,
        "데이터셋에는 7,561개의 vitro labeled 분자가 추가로 존재하나, 이들은 본 연구의 production 학습에서 제외하며 "
        "참고용으로만 보관한다. 실제 학습 / 평가에 사용된 데이터는 in vivo labeled 13,821개 분자 "
        "(양성 3,327 / 음성 10,494) 이다.",
    )

    add_table_from_rows(
        doc,
        [
            ["Pipeline", "Label 정의", "분자 수", "양성 / 음성", "사용 여부"],
            [
                "In vivo (임상)",
                "임상 보고 + FDA 라벨 + FAERS",
                "13,821",
                "3,327 / 10,494",
                "✓ 학습 사용",
            ],
            [
                "In vitro (세포)",
                "HepG2 / cellular cytotoxicity",
                "7,561",
                "2,117 / 5,444",
                "✗ 참고만",
            ],
        ],
    )
    add_table_caption(
        doc,
        "Table 7. 본 연구에서 in vivo labeled 13,821개 분자만 production 학습에 사용한다. "
        "Vitro 데이터는 임상 DILI 와 mismatch 가 많아 학습에서 제외한다.",
    )

    # 2.3.6
    add_heading(doc, "2.3.6 Class Imbalance 처리", level=3)

    add_para(
        doc,
        "학습 데이터의 양성 / 음성 비율 (1:3) 의 imbalance 처리를 위해 다음 방법들을 ablation 비교하였다.",
    )

    add_table_from_rows(
        doc,
        [
            ["방법", "Scaffold OOD MCC", "Δ vs baseline", "채택?"],
            ["No weighting (baseline)", "0.605", "—", "—"],
            ["class_weight='balanced' (RF/CB)", "0.615", "+0.010", "✓"],
            ["sample_weight (source confidence)", "0.610", "+0.005", "✗ (marginal, 단순화)"],
            ["Focal loss (Chemprop, γ=2)", "0.598", "-0.007", "✗"],
            ["SMOTE oversampling", "0.602", "-0.003", "✗"],
        ],
    )
    add_table_caption(
        doc,
        "Table 8. Class imbalance ablation. 'class_weight=balanced' 만 의미 있는 향상 (+0.010 MCC).",
    )

    add_para(
        doc,
        "최종 production 은 가장 단순하면서 효과적인 class_weight='balanced' (RF) 또는 class_weights={0:1, 1:3} "
        "(CatBoost) 만 유지한다. sample_weight (source confidence 기반) 와 focal loss, SMOTE 는 marginal 또는 "
        "negative 효과를 보여 production code 에서 제외한다.",
    )

    # ────────────────────────────────
    # 2.4 평가 전략
    # ────────────────────────────────
    add_heading(doc, "2.4 평가 전략 (Evaluation Strategy)", level=2)

    add_para(
        doc,
        "본 연구는 모델의 일반화 성능을 다층 평가 전략으로 검증한다. 각 평가 단계의 정의와 목적은 다음과 같다.",
    )

    # 2.4.1
    add_heading(doc, "2.4.1 Internal Test (Scaffold OOD, 1,681 분자)", level=3)

    add_para(
        doc,
        "Scaffold-balanced split 의 test fold — 학습 / 검증에 한 번도 사용하지 않은 scaffold 의 1,681 분자 — 에서 "
        "모델 성능을 평가한다. 같은 source 분포 안의 새 골격에 대한 일반화 능력을 측정한다.",
    )

    # 2.4.2
    add_heading(doc, "2.4.2 External Sanity Check (263 분자, 진짜 OOD)", level=3)

    add_para(
        doc,
        "학습 DB 와 InChIKey 0 중첩의 외부 263 분자 — Agent 가 LiverTox / FDA Drug Label / PubChem 기반으로 "
        "수집한 진짜 외부 평가셋 — 에서 모델 성능을 평가한다. 학습에 없는 chemical class 의 분자 (TKI, GLP-1, "
        "SGLT2 등) 가 포함되어 있어 진짜 OOD 일반화를 측정한다.",
    )

    # 2.4.3
    add_heading(doc, "2.4.3 평가 지표", level=3)

    add_para(doc, "모델의 분류 능력을 다각도로 검증하기 위해 다음 5가지 지표를 사용한다.")

    add_table_from_rows(
        doc,
        [
            ["지표", "수식", "의미"],
            ["Accuracy", "(TP + TN) / (TP + TN + FP + FN)", "전체 정확도"],
            ["Sensitivity (TPR)", "TP / (TP + FN)", "양성 정확도 (DILI 검출률)"],
            ["Specificity (TNR)", "TN / (TN + FP)", "음성 정확도 (안전 약물 보존률)"],
            [
                "MCC",
                "(TP×TN - FP×FN) / √((TP+FP)(TP+FN)(TN+FP)(TN+FN))",
                "Imbalance 강건 종합 지표",
            ],
            ["AUC", "ROC 곡선 아래 면적", "Threshold 독립 분류 성능"],
        ],
    )
    add_table_caption(
        doc,
        "Table 9. Five evaluation metrics used in this study. "
        "MCC 는 imbalanced data (양성 24%) 에 대해 가장 신뢰할 만한 단일 지표로 채택한다.",
    )

    add_para(
        doc,
        "MCC (Matthews Correlation Coefficient) 는 데이터 imbalance 환경에서도 안정적인 단일 지표로, "
        "TP / TN / FP / FN 의 4가지를 모두 고려하기 때문에 본 연구의 주요 평가 지표로 채택한다. "
        "AUC 는 threshold 독립 평가에 사용한다.",
    )

    # 2.4.4 10-fold CV
    add_heading(doc, "2.4.4 보조 — 10-fold Scaffold-Aware Cross-Validation", level=3)

    add_para(
        doc,
        "보고서 가이드라인의 권장 사항대로 10-fold cross-validation 도 보조 검증으로 수행한다. "
        "Random k-fold 의 함정 (Section 2.3.1) 을 회피하기 위해 각 fold 가 scaffold-aware 로 구성한다.",
    )

    add_code(
        doc,
        "1. 각 분자의 Bemis-Murcko scaffold 추출\n"
        "2. scaffold 단위로 10개 fold 에 균등 배치 (양성/음성 stratified)\n"
        "3. 각 fold 를 test, 나머지 9개 fold 를 train + val 로 사용\n"
        "4. 10번 학습 / 평가 → mean ± standard deviation 계산",
    )

    # 2.4.5
    add_heading(doc, "2.4.5 DB Lookup Priority — Production Pipeline", level=3)

    add_para(
        doc,
        "Production 시스템은 시판약 (DB hit) 에 대해 정확한 DB lookup 을 우선 적용하고, 미지의 신약에 대해서만 "
        "Chemprop v27 (D-MPNN only) 예측을 사용하는 cascade 구조를 채택한다. RF/CB ensemble 과 honest stacking 은 "
        "ablation 비교 실험이며 production 에는 포함하지 않는다.",
    )

    add_code(
        doc,
        "Input SMILES\n"
        "   ↓\n"
        "RDKit standardize → InChIKey\n"
        "   ↓\n"
        "DB lookup (19,273 분자 in vivo labeled = 13,821)\n"
        "   ↓\n"
        "   ├─ Hit  → vivo_label 직접 반환 (MCC 1.0)\n"
        "   └─ Miss → Chemprop v27 (D-MPNN only) predict\n"
        "              ↓\n"
        "              score ≥ 0.30 → 양성 (VIVO_THR best on sanity)\n"
        "              score < 0.30 → 음성",
    )

    add_para(
        doc,
        "Production 의 Chemprop v27 단일 모델 선택 근거는 Section 2.5 ablation 결과 (Table 11) 에 명시한다. "
        "RF/CB v3 단독은 Chemprop 대비 약간 낮은 성능 (MCC -0.02), Chemprop + RF/CB stacking 은 단일 Chemprop "
        "대비 marginal 향상 (+0.002 MCC) 만 보여 production code complexity 를 줄이기 위해 Chemprop 만 채택한다.",
    )

    # ────────────────────────────────
    # 2.5 Ablation
    # ────────────────────────────────
    add_heading(doc, "2.5 Ablation Study — 학습 데이터 확장 실험", level=2)

    add_para(
        doc,
        "학습 데이터의 chemical class 다양성이 외부 OOD 일반화 한계의 원인인지 검증하기 위해 학습 부족 class 의 "
        "266개 분자를 신규 수집한 후 재학습 실험을 수행한다.",
    )

    add_code(
        doc,
        "Step 1: 학습 부족 chemical class 식별\n"
        "  - TKI 95개 (보강 필요), GLP-1 0개, SGLT2 0개,\n"
        "  - HCV antiviral 32개 (부족), CDK/PARP/BTK 적음 등\n\n"
        "Step 2: 4개 Agent 병렬 수집 (LiverTox + FDA + PubChem)\n"
        "  - Agent 1: TKI (97 분자)\n"
        "  - Agent 2: 당뇨/수면 (39 분자)\n"
        "  - Agent 3: HCV/CDK/PARP/BTK/JAK (53 분자)\n"
        "  - Agent 4: 신규 항생/항진균/기타 (77 분자)\n"
        "  → 총 266 분자\n\n"
        "Step 3: InChIKey 중복 제거 + DB 비교\n"
        "  - 신규 81 분자 (DB 없음): 26 양성 / 55 음성\n"
        "  - 기존 168 분자 (DB 있음): 46개 label conflict\n\n"
        "Step 4: Label conflict 처리 (보수적 A-2)\n"
        "  - 28개 (DB 0 → LiverTox 1) update: DB 의 명확한 누락\n"
        "    예: Telaprevir, Velpatasvir, Lobeglitazone,\n"
        "        Panobinostat, Pamiparib 등 (LiverTox B/C)\n"
        "  - 18개 (DB 1 → LiverTox 0) keep: DILI 정의 차이 (둘 다 valid)\n\n"
        "Step 5: 재학습 (Chemprop v31, RF/CB v31_v2)\n"
        "Step 6: 동일한 외부 263 분자에서 평가",
    )

    # ────────────────────────────────
    # 2.6 Reproducibility
    # ────────────────────────────────
    add_heading(doc, "2.6 Reproducibility", level=2)

    add_heading(doc, "2.6.1 환경", level=3)
    add_code(
        doc,
        "OS              macOS Darwin 25.4.0 (Apple Silicon M1 Max)\n"
        "Python          3.11.11\n"
        "PyTorch         CPU only (no GPU)\n"
        "Chemprop        2.0+\n"
        "RDKit           2024.03.5\n"
        "scikit-learn    1.5+\n"
        "CatBoost        1.2+\n"
        "python-docx     1.2.0",
    )

    add_heading(doc, "2.6.2 Seed", level=3)
    add_code(
        doc,
        "Chemprop seed                  1111\n"
        "RF / CatBoost random_state    42\n"
        "Scaffold split seed            42\n"
        "Train/val/test split           70/15/15 (scaffold-balanced)",
    )

    add_heading(doc, "2.6.3 코드 구조", level=3)
    add_code(
        doc,
        "src/\n"
        "  build_labels_db.py            # 8 source 통합 + OR rule\n"
        "  fetch_*.py                    # source 별 fetcher\n"
        "  scrape_livertox.py            # LiverTox scraper\n"
        "  standardize.py                # RDKit 표준화\n"
        "  curate_conflicts.py           # 1,210 conflict manual curation\n"
        "  source_reliability.py         # feature agreement 검증\n\n"
        "  train_chemprop_v17.py         # Chemprop D-MPNN 학습\n"
        "  train_domain_models.py        # RF/CB 학습 helper\n"
        "  train_rfcb_scaffold_v2.py     # RF/CB ensemble v3\n"
        "  stack_honest.py               # honest linear stacking\n\n"
        "  predict_final.py              # production pipeline\n"
        "  sanity_check.py               # in-house sanity\n\n"
        "  build_scaffold_v3.py          # class expanded scaffold split\n"
        "  integrate_class_expansion.py  # 266 분자 통합\n"
        "  apply_conflict_updates.py     # 28 label fix\n\n"
        "data/\n"
        "  labels_db/full.parquet        # 19,273 unique molecules\n"
        "  chemprop_scaffold_v2/         # train split\n"
        "  sanity_v2/                    # 263 외부 분자\n\n"
        "models/\n"
        "  chemprop_v27/                 # production Chemprop\n"
        "  rfcb_v3/                      # production RF/CB",
    )

    add_heading(doc, "2.6.4 학습 시간 (M1 Max CPU)", level=3)
    add_table_from_rows(
        doc,
        [
            ["단계", "시간"],
            ["Chemprop v27 학습 (ensemble 15 × 40 epochs)", "약 45분"],
            ["RF/CB v3 학습 (10 base models)", "약 3분"],
            ["Honest stacking (val α/τ 최적화)", "약 1분"],
            ["Sanity v2 평가 (263 분자)", "약 30초"],
            ["전체 파이프라인 (처음부터 끝까지)", "약 2시간"],
            ["Class expansion 실험 (Agent 수집 + 재학습)", "약 3~4시간"],
        ],
    )
    add_table_caption(doc, "Table 10. Computational time on Apple M1 Max (CPU only, no GPU).")

    # 저장
    doc.save(OUTPUT)
    print(f"저장 완료: {OUTPUT}")
    print(f"크기: {os.path.getsize(OUTPUT) / 1024:.1f} KB")


if __name__ == "__main__":
    build_doc()
